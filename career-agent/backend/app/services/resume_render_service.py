from __future__ import annotations

import html
import re
from pathlib import Path
from typing import Any

from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


class ResumeRenderService:
    """Render structured resume data into both DOCX and PDF formats."""

    advisory_pattern = re.compile(
        r"(建议补充[^。；;]*[。；;]?|建议[^。；;]*[。；;]?|需要[^。；;]*[。；;]?|应当[^。；;]*[。；;]?|"
        r"recommend[^.;]*[.;]?|suggest[^.;]*[.;]?|should[^.;]*[.;]?)",
        re.I,
    )

    @classmethod
    def _clean_advisory_text(cls, text: str) -> str:
        cleaned = cls.advisory_pattern.sub("", str(text or ""))
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
        return cleaned

    def render_word(self, *, resume_document: dict[str, Any], output_path: Path) -> Path:
        try:
            from docx import Document  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.oxml.ns import qn  # type: ignore
            from docx.shared import Cm, Pt  # type: ignore
        except Exception as exc:  # pragma: no cover - dependency guard
            raise RuntimeError(f"word export dependency is missing: {exc}") from exc

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.25)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Microsoft YaHei"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal_style.font.size = Pt(9.5)
        normal_style.paragraph_format.space_after = Pt(2)
        normal_style.paragraph_format.line_spacing = 1.08

        name = str(resume_document.get("name") or "").strip()
        target_role = str(resume_document.get("target_role") or "").strip()
        header_title = self._join_non_empty("", name, target_role, sep=" - ") or "个人简历"
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(header_title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = "Microsoft YaHei"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        title_paragraph.paragraph_format.space_after = Pt(3)

        meta_line = self._build_meta_line(resume_document)
        if meta_line:
            meta = document.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta.add_run(meta_line)
            run.font.size = Pt(8.8)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            meta.paragraph_format.space_after = Pt(7)

        summary = self._clean_advisory_text(str(resume_document.get("summary") or "").strip())
        if summary:
            self._add_word_heading(document, "个人简介")
            self._add_word_body(document, summary)

        education_experience = self._clean_advisory_text(str(resume_document.get("education_experience") or "").strip())
        if education_experience:
            self._add_word_heading(document, "教育经历")
            for index, line in enumerate(self._split_lines(education_experience)):
                self._add_word_body(document, line, bold=index == 0)

        skills = self._normalize_str_list(resume_document.get("skills"))
        if skills:
            self._add_word_heading(document, "核心技能")
            self._add_word_body(document, "、".join(skills))

        certificates = self._normalize_str_list(resume_document.get("certificates"))
        if certificates:
            self._add_word_heading(document, "证书与资质")
            for item in certificates:
                self._add_word_bullet(document, item)

        projects = resume_document.get("projects") or []
        if projects:
            self._add_word_heading(document, "项目经历")
            for project in projects:
                title_line = self._join_non_empty(
                    "",
                    project.get("name"),
                    self._join_non_empty("", project.get("role"), project.get("duration"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    self._add_word_body(document, title_line, bold=True)
                for bullet in self._split_bullet_text(project.get("rewrite")):
                    self._add_word_bullet(document, bullet)

        internships = resume_document.get("internships") or []
        if internships:
            self._add_word_heading(document, "实习经历")
            for internship in internships:
                title_line = self._join_non_empty(
                    "",
                    internship.get("company"),
                    self._join_non_empty("", internship.get("position"), internship.get("duration"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    self._add_word_body(document, title_line, bold=True)
                for bullet in self._split_bullet_text(internship.get("rewrite")):
                    self._add_word_bullet(document, bullet)

        competitions = resume_document.get("competitions") or []
        if competitions:
            self._add_word_heading(document, "竞赛经历")
            for item in competitions:
                title_line = self._join_non_empty(
                    "",
                    item.get("name"),
                    self._join_non_empty("", item.get("award"), item.get("level"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    self._add_word_body(document, title_line, bold=True)
                for bullet in self._split_bullet_text(item.get("description")):
                    self._add_word_bullet(document, bullet)

        campus_experiences = resume_document.get("campus_experiences") or []
        if campus_experiences:
            self._add_word_heading(document, "校园经历")
            for item in campus_experiences:
                title_line = self._join_non_empty(
                    "",
                    item.get("title"),
                    self._join_non_empty("", item.get("role"), item.get("duration"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    self._add_word_body(document, title_line, bold=True)
                for bullet in self._split_bullet_text(item.get("description")):
                    self._add_word_bullet(document, bullet)

        document.save(str(output_path))
        return output_path

    def render_pdf(self, *, resume_document: dict[str, Any], output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)

        registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        if "ResumeTitle" not in styles:
            styles.add(ParagraphStyle(name="ResumeTitle", fontName="STSong-Light", fontSize=17, leading=22, alignment=TA_CENTER))
        if "ResumeMeta" not in styles:
            styles.add(ParagraphStyle(name="ResumeMeta", fontName="STSong-Light", fontSize=8.8, leading=12, alignment=TA_CENTER))
        if "ResumeHeading" not in styles:
            styles.add(ParagraphStyle(name="ResumeHeading", fontName="STSong-Light", fontSize=11.2, leading=15, spaceBefore=4, spaceAfter=3))
        if "ResumeBody" not in styles:
            styles.add(ParagraphStyle(name="ResumeBody", fontName="STSong-Light", fontSize=9.4, leading=14))
        if "ResumeBullet" not in styles:
            styles.add(ParagraphStyle(name="ResumeBullet", fontName="STSong-Light", fontSize=9.4, leading=14, leftIndent=10))

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=15 * mm,
            rightMargin=15 * mm,
            topMargin=13 * mm,
            bottomMargin=12 * mm,
        )
        story: list[Any] = []

        name = str(resume_document.get("name") or "").strip()
        target_role = str(resume_document.get("target_role") or "").strip()
        header_title = self._join_non_empty("", name, target_role, sep=" - ") or "个人简历"
        story.append(Paragraph(self._escape(header_title), styles["ResumeTitle"]))

        meta_line = self._build_meta_line(resume_document)
        if meta_line:
            story.append(Paragraph(self._escape(meta_line), styles["ResumeMeta"]))
        story.append(Spacer(1, 5))

        summary = self._clean_advisory_text(str(resume_document.get("summary") or "").strip())
        if summary:
            self._append_pdf_heading(story, styles["ResumeHeading"], "个人简介")
            story.append(Paragraph(self._escape(summary), styles["ResumeBody"]))

        education_experience = self._clean_advisory_text(str(resume_document.get("education_experience") or "").strip())
        if education_experience:
            self._append_pdf_heading(story, styles["ResumeHeading"], "教育经历")
            for line in self._split_lines(education_experience):
                story.append(Paragraph(self._escape(line), styles["ResumeBody"]))

        skills = self._normalize_str_list(resume_document.get("skills"))
        if skills:
            self._append_pdf_heading(story, styles["ResumeHeading"], "核心技能")
            story.append(Paragraph(self._escape("、".join(skills)), styles["ResumeBody"]))

        certificates = self._normalize_str_list(resume_document.get("certificates"))
        if certificates:
            self._append_pdf_heading(story, styles["ResumeHeading"], "证书与资质")
            for item in certificates:
                story.append(Paragraph(self._escape(f"• {item}"), styles["ResumeBullet"]))

        projects = resume_document.get("projects") or []
        if projects:
            self._append_pdf_heading(story, styles["ResumeHeading"], "项目经历")
            for project in projects:
                title_line = self._join_non_empty(
                    "",
                    project.get("name"),
                    self._join_non_empty("", project.get("role"), project.get("duration"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    story.append(Paragraph(self._escape(title_line), styles["ResumeBody"]))
                for bullet in self._split_bullet_text(project.get("rewrite")):
                    story.append(Paragraph(self._escape(f"• {bullet}"), styles["ResumeBullet"]))

        internships = resume_document.get("internships") or []
        if internships:
            self._append_pdf_heading(story, styles["ResumeHeading"], "实习经历")
            for internship in internships:
                title_line = self._join_non_empty(
                    "",
                    internship.get("company"),
                    self._join_non_empty("", internship.get("position"), internship.get("duration"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    story.append(Paragraph(self._escape(title_line), styles["ResumeBody"]))
                for bullet in self._split_bullet_text(internship.get("rewrite")):
                    story.append(Paragraph(self._escape(f"• {bullet}"), styles["ResumeBullet"]))

        competitions = resume_document.get("competitions") or []
        if competitions:
            self._append_pdf_heading(story, styles["ResumeHeading"], "竞赛经历")
            for item in competitions:
                title_line = self._join_non_empty(
                    "",
                    item.get("name"),
                    self._join_non_empty("", item.get("award"), item.get("level"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    story.append(Paragraph(self._escape(title_line), styles["ResumeBody"]))
                for bullet in self._split_bullet_text(item.get("description")):
                    story.append(Paragraph(self._escape(f"• {bullet}"), styles["ResumeBullet"]))

        campus_experiences = resume_document.get("campus_experiences") or []
        if campus_experiences:
            self._append_pdf_heading(story, styles["ResumeHeading"], "校园经历")
            for item in campus_experiences:
                title_line = self._join_non_empty(
                    "",
                    item.get("title"),
                    self._join_non_empty("", item.get("role"), item.get("duration"), sep=" | "),
                    sep=" - ",
                )
                if title_line:
                    story.append(Paragraph(self._escape(title_line), styles["ResumeBody"]))
                for bullet in self._split_bullet_text(item.get("description")):
                    story.append(Paragraph(self._escape(f"• {bullet}"), styles["ResumeBullet"]))

        doc.build(story)
        return output_path

    @staticmethod
    def _add_word_heading(document: Any, text: str) -> None:
        from docx.shared import Pt  # type: ignore

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(11)

    @staticmethod
    def _add_word_body(document: Any, text: str, *, bold: bool = False) -> None:
        from docx.shared import Pt  # type: ignore

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1.5)
        run = paragraph.add_run(str(text or "").strip())
        run.bold = bold
        run.font.size = Pt(9.5)

    @staticmethod
    def _add_word_bullet(document: Any, text: str) -> None:
        from docx.shared import Pt  # type: ignore

        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(str(text or "").strip())
        run.font.size = Pt(9.5)

    @staticmethod
    def _append_pdf_heading(story: list[Any], style: ParagraphStyle, text: str) -> None:
        story.append(Spacer(1, 4))
        story.append(Paragraph(ResumeRenderService._escape(text), style))

    def _build_meta_line(self, resume_document: dict[str, Any]) -> str:
        school_major = self._join_non_empty("", resume_document.get("college"), resume_document.get("major"), sep=" · ")
        links = self._normalize_str_list(resume_document.get("links"))
        github = str(resume_document.get("github") or "").strip()
        meta = [
            resume_document.get("phone"),
            resume_document.get("email"),
            github,
            *[item for item in links if item != github],
            resume_document.get("target_city"),
            school_major,
            resume_document.get("grade"),
        ]
        return self._join_non_empty("", *meta, sep=" | ")

    @classmethod
    def _split_bullet_text(cls, value: Any) -> list[str]:
        text = cls._clean_advisory_text(str(value or "").strip())
        if not text:
            return []
        parts = [item.strip(" 。；;") for item in re.split(r"[\n；;]+", text) if item.strip(" 。；;")]
        return parts or [text]

    @staticmethod
    def _split_lines(value: str) -> list[str]:
        return [line.strip() for line in str(value or "").splitlines() if line.strip()]

    @staticmethod
    def _normalize_str_list(value: Any) -> list[str]:
        if not value:
            return []
        if isinstance(value, str):
            value = re.split(r"[,，/、;；\n]+", value)
        rows: list[str] = []
        for item in value:
            text = str(item or "").strip()
            if text:
                rows.append(text)
        return rows

    @staticmethod
    def _join_non_empty(prefix: str, *values: Any, sep: str = "") -> str:
        parts = [str(item).strip() for item in values if str(item or "").strip()]
        if not parts:
            return ""
        body = sep.join(parts) if sep else "".join(parts)
        return f"{prefix}{body}" if prefix else body

    @staticmethod
    def _escape(value: str) -> str:
        return html.escape(str(value or "")).replace("\n", "<br/>")
