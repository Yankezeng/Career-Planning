from __future__ import annotations

import json
import logging
import re
import socket
import ssl
import time
from hashlib import sha256
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Callable
from urllib.error import URLError
from urllib.request import Request, urlopen
from uuid import uuid4

from sqlalchemy.orm import Session, joinedload

from app.core.config import get_settings
from app.models.auth import User
from app.models.student import Student, StudentAttachment, StudentResume, StudentResumeVersion
from app.services.agent.code_agent.code_agent import CodeAgent
from app.services.agent.common.agent_llm_profiles import AgentLLMProfile, get_agent_llm_profile
from app.services.llm_contracts import LLMCallError
from app.services.resume_optimizer_service import ResumeOptimizerService
from app.services.resume_parser_service import ResumeParserService
from app.services.resume_render_service import ResumeRenderService
from app.services.structured_llm_service import get_structured_llm_service_for_profile
from app.utils.upload_paths import resolve_upload_reference, upload_path_to_url


class DocumentBuildSpecTimeoutError(RuntimeError):
    def __init__(self, message: str, *, tool_steps: list[dict[str, Any]] | None = None):
        super().__init__(message)
        self.tool_steps = tool_steps or []


ResumeJobPhaseCallback = Callable[[str, str], None]


class FileAgent:
    FILE_TASKS = {
        "parse_file",
        "optimize_resume",
        "generate_report",
        "generate_document",
        "export_optimized_resume",
        "generate_chart",
        "generate_image",
    }
    DOCUMENT_BUILD_SPEC_REQUIRED_FIELDS = (
        "document_type",
        "target_role",
        "enterprise_alignment",
        "career_planning",
        "revision_suggestions",
        "optimized_resume_document",
        "document_style",
        "sections",
        "supervisor_checklist",
    )

    def __init__(self, db: Session, llm_profile: AgentLLMProfile | None = None):
        self.db = db
        self.settings = get_settings()
        self.llm_profile = llm_profile or get_agent_llm_profile("file_agent")
        self.resume_parser = ResumeParserService(
            text_model=self.llm_profile.module_name,
            vision_model=self.llm_profile.module_name,
            api_key=self.llm_profile.api_key,
            base_url=self.llm_profile.base_url,
        )
        self.resume_optimizer = ResumeOptimizerService(db, resume_parser=self.resume_parser)
        from app.services.report_service_v2 import ReportService

        self.report_service = ReportService(
            db,
            llm_service=get_structured_llm_service_for_profile(
                api_key=self.llm_profile.api_key,
                base_url=self.llm_profile.base_url,
                module_name=self.llm_profile.module_name,
            ),
        )

    def detect_task(self, message: str, selected_skill: str = "") -> str:
        text = str(message or "").strip().lower()
        compact = re.sub(r"\s+", "", text)
        skill = str(selected_skill or "").strip().lower()

        if any(token in compact for token in ["解析", "识别", "提取", "分析", "查看", "读取", "parse", "extract"]) and any(
            token in compact for token in ["图片", "pdf", "word", "doc", "docx", "文件", "简历", "image", "file", "resume"]
        ):
            return "parse_file"
        if any(token in compact for token in ["优化简历", "简历优化", "简历修改", "optimize resume"]) or (
            "简历" in compact and any(token in compact for token in ["优化", "修改", "润色", "改写", "rewrite", "optimize"])
        ):
            return "optimize_resume"
        if "resume-workbench" in skill and any(token in compact for token in ["优化", "optimize", "rewrite"]):
            return "optimize_resume"
        if any(token in compact for token in ["报告", "report"]) and any(token in compact for token in ["生成", "导出", "下载", "generate", "export", "download"]):
            return "generate_report"
        negative_chart_intent = any(token in compact for token in ["不要生成图表", "不生成图表", "无需图表", "不要图表", "nochart"])
        explicit_chart_action = any(token in compact for token in ["生成", "创建", "绘制", "画", "做", "generate", "create", "draw", "chart"])
        chart_kind = any(token in compact for token in ["柱状图", "折线图", "饼图"])
        if not negative_chart_intent and any(token in compact for token in ["图表", "柱状图", "折线图", "饼图", "chart"]) and (explicit_chart_action or chart_kind):
            return "generate_chart"
        if any(token in compact for token in ["生成图片", "创建图片", "海报", "配图", "generateimage", "createimage"]):
            return "generate_image"
        if any(token in compact for token in ["文档", "word", "docx", "document"]) and any(
            token in compact for token in ["生成", "创建", "下载", "generate", "create", "download"]
        ):
            return "generate_document"
        return ""

    def execute(
        self,
        *,
        user: User,
        message: str,
        task_type: str,
        session_state: dict[str, Any] | None = None,
        context_binding: dict[str, Any] | None = None,
        client_state: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        session_state = session_state or {}
        context_binding = context_binding or {}
        client_state = client_state or {}

        student = self._student_by_user(user.id)
        if not student:
            return self._needs_input(
                task_type=task_type,
                question="当前账号未关联学生档案，暂时无法执行文件任务。",
                missing_fields=["student_profile"],
            )
        if task_type == "parse_file":
            return self._parse_file(student, message, session_state, context_binding, client_state)
        if task_type == "optimize_resume":
            return self._optimize_resume(user, student, message, session_state, context_binding, client_state)
        if task_type == "generate_report":
            return self._generate_report(student)
        if task_type == "export_optimized_resume":
            return self._export_optimized_resume_from_context(
                student,
                session_state=session_state,
                context_binding=context_binding,
                client_state=client_state,
            )
        if task_type == "generate_document":
            return self._generate_document(
                student,
                message,
                session_state=session_state,
                context_binding=context_binding,
                client_state=client_state,
            )
        if task_type == "generate_chart":
            return self._generate_chart(student, message)
        if task_type == "generate_image":
            return self._generate_image(student, message)

        return self._needs_input(task_type=task_type, question="暂不支持该文件任务。", missing_fields=["task_type"])

    def _parse_file(
        self,
        student: Student,
        message: str,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        attachment_resolution = self._resolve_attachment_context(
            student.id,
            session_state,
            context_binding,
            client_state,
        )
        attachment = attachment_resolution.get("attachment")
        if not attachment:
            return self._needs_input(
                task_type="parse_file",
                question=str(attachment_resolution.get("question") or "请先上传图片、Word 或 PDF 文件，然后我再继续解析。"),
                missing_fields=["attachment_id"],
                invalid_fields=list(attachment_resolution.get("invalid_fields") or []),
                context_patch=attachment_resolution.get("context_patch") if isinstance(attachment_resolution.get("context_patch"), dict) else None,
            )

        file_path = self._resolve_attachment_abs_path(attachment)
        if not file_path:
            return self._needs_input(
                task_type="parse_file",
                question="已找到附件记录，但对应文件不存在。请重新上传后再试。",
                missing_fields=["attachment_id"],
            )
        extraction_report, extract_steps, extract_failure = self._extract_file_with_code_agent(
            user=None,
            student=student,
            attachment=attachment,
            file_path=file_path,
            message=message,
        )
        if extract_failure:
            return self._failed(
                task_type="parse_file",
                failure_reason=extract_failure,
                tool_steps=extract_steps,
            )
        if self._is_extraction_report_low_quality(extraction_report, require_text=False):
            return self._failed(
                task_type="parse_file",
                failure_reason="附件解析质量不足，已停止文件解析任务。",
                tool_steps=[*extract_steps, {"tool": "parse_file", "status": "failed", "text": "failed: parse quality low"}],
            )
        parsed = self._build_parsed_resume_from_extraction_report(attachment=attachment, extraction_report=extraction_report)
        return self._build_parse_success_result(
            attachment=attachment,
            parsed=parsed,
            parse_source="attachment",
            summary=f"已解析文件《{attachment.file_name}》。",
        )

    def _build_parse_success_result(
        self,
        *,
        attachment: StudentAttachment | None,
        parsed: dict[str, Any],
        parse_source: str,
        summary: str,
    ) -> dict[str, Any]:
        attachment_id = int(getattr(attachment, "id", 0) or 0)
        attachment_name = str(getattr(attachment, "file_name", "") or "")
        attachment_type = str(getattr(attachment, "file_type", "") or "")
        tool_output = self._tool_output(
            tool="parse_file",
            title="文件解析",
            summary=summary,
            data={
                "attachment_id": attachment_id,
                "attachment_name": attachment_name,
                "attachment_type": attachment_type,
                "parse_source": parse_source,
                "parser_success": bool(parsed.get("parser_success", True)),
                "parsed": parsed,
            },
            next_actions=["继续上传文件解析", "继续生成可下载文档"],
            card_type="resume_card",
            context_patch=self._build_file_parse_context_patch(attachment=attachment, status="", reason=""),
        )
        return self._success(
            task_type="parse_file",
            reply=summary,
            tool_steps=[{"tool": "parse_file", "status": "done", "text": "done: 文件解析"}],
            artifacts=[],
            context_patch=tool_output.get("context_patch") or {},
            tool_outputs=[tool_output],
        )

    def _optimize_resume(
        self,
        user: User,
        student: Student,
        message: str,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        attachment_resolution = self._resolve_attachment_context(
            student.id,
            session_state,
            context_binding,
            client_state,
        )
        attachment = attachment_resolution.get("attachment")
        if not attachment:
            return self._needs_input(
                task_type="optimize_resume",
                question=str(attachment_resolution.get("question") or "请先上传简历附件（图片、Word 或 PDF），然后我再继续优化并导出。"),
                missing_fields=["attachment_id"],
                invalid_fields=list(attachment_resolution.get("invalid_fields") or []),
                context_patch=attachment_resolution.get("context_patch") if isinstance(attachment_resolution.get("context_patch"), dict) else None,
            )

        file_path = self._resolve_attachment_abs_path(attachment)
        if not file_path:
            return self._failed(
                task_type="optimize_resume",
                failure_reason="已找到附件记录，但本地文件不存在。请重新上传简历附件后重试。",
                tool_steps=[{"tool": "local_parse", "status": "failed", "text": "failed: attachment file not found"}],
            )

        optimization_options = self._resolve_resume_optimization_options(
            session_state=session_state,
            context_binding=context_binding,
            client_state=client_state,
        )
        return self._build_resume_optimization_background_result(
            student=student,
            attachment=attachment,
            message=message,
            optimization_options=optimization_options,
        )
        if self._is_extraction_report_low_quality(extraction_report, require_text=True):
            return self._failed(
                task_type="optimize_resume",
                failure_reason="CodeAgent 本地解析出的简历正文为空或质量不足，已停止优化。请确认附件内容可复制或换成可解析的 Word/PDF。",
                tool_steps=[*extract_steps, {"tool": "code_file_extract", "status": "failed", "text": "failed: extracted resume text is empty"}],
            )
        parsed = self._build_parsed_resume_from_extraction_report(attachment=attachment, extraction_report=extraction_report)
        try:
            return self._complete_resume_optimization(
                user=user,
                student=student,
                attachment=attachment,
                parsed=parsed,
                extraction_report=extraction_report,
                extract_steps=extract_steps,
                message=message,
                optimization_options=optimization_options,
            )
        except TimeoutError as exc:
            return self._build_resume_optimization_background_result(
                student=student,
                attachment=attachment,
                parsed=parsed,
                extraction_report=extraction_report,
                extract_steps=extract_steps,
                message=message,
                optimization_options=optimization_options,
                timeout_error=str(exc),
            )
        except (LLMCallError, URLError, json.JSONDecodeError, ValueError, KeyError, IndexError, TypeError) as exc:
            return self._failed(
                task_type="optimize_resume",
                failure_reason=f"简历优化服务处理失败：{str(exc)}",
                tool_steps=[*extract_steps, {"tool": "qwen_resume_optimization", "status": "failed", "text": f"failed: {str(exc)}"}],
            )

    def _complete_resume_optimization(
        self,
        *,
        user: User,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        extraction_report: dict[str, Any],
        extract_steps: list[dict[str, Any]],
        message: str,
        optimization_options: dict[str, Any],
        timeout_seconds: int | None = None,
        max_retries: int | None = None,
        optimization_result: dict[str, Any] | None = None,
        optimization_tool_steps: list[dict[str, Any]] | None = None,
        cached_export_artifact: dict[str, Any] | None = None,
        phase_callback: ResumeJobPhaseCallback | None = None,
    ) -> dict[str, Any]:
        if optimization_result is None:
            self._emit_resume_job_phase(phase_callback, "optimizing", "正在按目标岗位优化简历...")
            optimization_result, optimization_tool_steps = self._build_resume_optimization_result(
                student=student,
                attachment=attachment,
                parsed=parsed,
                extraction_report=extraction_report,
                message=message,
                optimization_options=optimization_options,
                timeout_seconds=timeout_seconds,
                max_retries=max_retries,
            )
        else:
            optimization_tool_steps = optimization_tool_steps or []

        wants_export = self._wants_resume_file_output(message)
        if not wants_export:
            summary = "已基于本地解析内容完成简历优化建议；本次未检测到导出 Word/DOCX/PDF 的明确意图，因此不生成文件。"
            reply = self._format_resume_optimization_reply(optimization_result, fallback_summary=summary)
            tool_output = self._tool_output(
                tool="optimize_resume",
                title="简历优化",
                summary=summary,
                data={
                    "attachment_id": attachment.id,
                    "target_role": optimization_result.get("target_role") or optimization_options.get("target_role") or "",
                    "resume_optimization_result": optimization_result,
                    "file_extraction_report": extraction_report,
                    "parsed_resume": parsed,
                    "export_requested": False,
                    "artifacts": [],
                },
                next_actions=["如需文件，请发送：优化并导出 Word", "继续岗位匹配分析"],
                card_type="resume_card",
                context_patch={
                    "context_binding": {
                        "attachment_id": attachment.id,
                        "attachment": {
                            "id": attachment.id,
                            "file_name": attachment.file_name,
                            "file_type": attachment.file_type,
                        },
                        "resume_optimization_result": optimization_result,
                        "last_resume_optimization": {
                            "attachment_id": attachment.id,
                            "source_file_name": attachment.file_name,
                            "target_role": optimization_result.get("target_role"),
                            "optimization_result": optimization_result,
                            "parsed_resume": parsed,
                            "file_extraction_report": extraction_report,
                            "created_at": datetime.now().isoformat(timespec="seconds"),
                        },
                    }
                },
            )
            result = self._success(
                task_type="optimize_resume",
                reply=reply,
                tool_steps=[*extract_steps, *optimization_tool_steps],
                artifacts=[],
                context_patch=tool_output.get("context_patch") or {},
                tool_outputs=[tool_output],
            )
            result["agent_flow_patch"] = [
                {"agent": "file", "action": "resolve_resume_attachment"},
                {"agent": "code", "action": "file_extract_request"},
                {"agent": "file", "action": "qwen_resume_optimization"},
                {"agent": "supervisor", "action": "return_text_suggestions"},
            ]
            return result

        if cached_export_artifact:
            summary = "已复用本次简历优化缓存中的 Word 文件。"
            tool_output = self._tool_output(
                tool="optimize_resume",
                title="简历优化",
                summary=summary,
                data={
                    "attachment_id": attachment.id,
                    "target_role": optimization_result.get("target_role") or optimization_options.get("target_role") or "",
                    "resume_optimization_result": optimization_result,
                    "file_extraction_report": extraction_report,
                    "parsed_resume": parsed,
                    "artifacts": [cached_export_artifact],
                    "export_requested": True,
                    "cache_hit": True,
                },
                next_actions=["继续生成职业规划报告", "继续岗位匹配分析"],
                card_type="resume_card",
                context_patch={
                    "context_binding": {
                        "attachment_id": attachment.id,
                        "attachment": {
                            "id": attachment.id,
                            "file_name": attachment.file_name,
                            "file_type": attachment.file_type,
                        },
                        "last_resume_optimization": {
                            "attachment_id": attachment.id,
                            "source_file_name": attachment.file_name,
                            "target_role": optimization_result.get("target_role"),
                            "optimization_result": optimization_result,
                            "parsed_resume": parsed,
                            "file_extraction_report": extraction_report,
                            "created_at": datetime.now().isoformat(timespec="seconds"),
                            "artifact_id": cached_export_artifact.get("id"),
                        },
                    }
                },
            )
            result = self._success(
                task_type="optimize_resume",
                reply=summary,
                tool_steps=[
                    *extract_steps,
                    *optimization_tool_steps,
                    {"tool": "code_document_render_cache", "status": "done", "text": "done: reused cached DOCX artifact"},
                ],
                artifacts=[cached_export_artifact],
                context_patch=tool_output.get("context_patch") or {},
                tool_outputs=[tool_output],
            )
            result["agent_flow_patch"] = [
                {"agent": "file", "action": "resolve_resume_attachment"},
                {"agent": "code", "action": "file_extract_cache"},
                {"agent": "file", "action": "qwen_resume_optimization_cache"},
                {"agent": "file", "action": "reuse_docx_artifact"},
                {"agent": "supervisor", "action": "check_document_result"},
            ]
            return result

        self._emit_resume_job_phase(phase_callback, "rendering_word", "正在生成 Word 文件...")
        spec = self._build_document_spec_from_resume_optimization(
            student=student,
            attachment=attachment,
            parsed=parsed,
            optimization_result=optimization_result,
            optimization_options=optimization_options,
        )
        output_name = self._sanitize_display_filename(
            file_name=f"{Path(attachment.file_name).stem}-career-optimized.docx",
            default_name="career-optimized-resume.docx",
        )
        document_render_request = {
            "spec": spec,
            "output_format": "docx",
            "output_name": output_name,
            "source_attachment_id": attachment.id,
            "student_id": student.id,
            "render_style": "career_enterprise_polished",
        }

        code_result = CodeAgent(self.db).execute(
            user=user,
            message=self._build_code_agent_document_render_prompt(document_render_request=document_render_request),
            selected_skill="code-agent",
            context_binding={"document_render_request": document_render_request},
            session_state={},
            client_state={},
        )
        if str(code_result.get("status") or "") != "success":
            failure_reason = str((code_result.get("code_task") or {}).get("failure_reason") or code_result.get("question") or "CodeAgent 文档渲染失败。")
            return self._failed(
                task_type="optimize_resume",
                failure_reason=failure_reason,
                tool_steps=[
                    *extract_steps,
                    *optimization_tool_steps,
                    {"tool": "code_document_render", "status": "failed", "text": "failed: CodeAgent document_render"},
                ],
            )

        render_report = self._extract_document_render_report(code_result)
        output_path = Path(str(render_report.get("output_path") or ""))
        if not output_path.exists():
            return self._failed(
                task_type="optimize_resume",
                failure_reason="CodeAgent 返回成功，但未产生可登记的 DOCX 文件。",
                tool_steps=[
                    *extract_steps,
                    *optimization_tool_steps,
                    {"tool": "code_document_render", "status": "failed", "text": "failed: missing rendered DOCX"},
                ],
            )

        self._emit_resume_job_phase(phase_callback, "registering_artifact", "正在登记 Word 下载文件...")
        artifact = self._register_generated_attachment(
            student_id=student.id,
            abs_path=output_path,
            file_name=output_name,
            description="agent generated career enterprise polished resume docx",
        )
        summary = "已按企业匹配与职业规划方向完成简历优化，并生成可下载 DOCX 文档。"
        tool_output = self._tool_output(
            tool="optimize_resume",
            title="简历优化",
            summary=summary,
            data={
                "attachment_id": attachment.id,
                "target_role": spec.get("target_role") or "",
                "resume_optimization_result": optimization_result,
                "document_build_spec": spec,
                "document_render_request": document_render_request,
                "code_render_report": render_report,
                "file_extraction_report": extraction_report,
                "parsed_resume": parsed,
                "artifacts": [artifact],
                "export_requested": True,
            },
            next_actions=["继续生成职业规划报告", "继续岗位匹配分析"],
            card_type="resume_card",
            context_patch={
                "context_binding": {
                    "attachment_id": attachment.id,
                    "document_render_request": document_render_request,
                    "attachment": {
                        "id": attachment.id,
                        "file_name": attachment.file_name,
                        "file_type": attachment.file_type,
                    },
                    "last_resume_optimization": {
                        "attachment_id": attachment.id,
                        "source_file_name": attachment.file_name,
                        "target_role": optimization_result.get("target_role"),
                        "optimization_result": optimization_result,
                        "parsed_resume": parsed,
                        "file_extraction_report": extraction_report,
                        "created_at": datetime.now().isoformat(timespec="seconds"),
                        "artifact_id": artifact.get("id"),
                    },
                }
            },
        )
        result = self._success(
            task_type="optimize_resume",
            reply=summary,
            tool_steps=[
                *extract_steps,
                *optimization_tool_steps,
                {"tool": "code_document_render", "status": "done", "text": "done: CodeAgent DOCX render"},
                {"tool": "register_artifact", "status": "done", "text": "done: registered DOCX artifact"},
            ],
            artifacts=[artifact],
            context_patch=tool_output.get("context_patch") or {},
            tool_outputs=[tool_output],
        )
        result["agent_flow_patch"] = [
            {"agent": "file", "action": "resolve_resume_attachment"},
            {"agent": "code", "action": "file_extract_request"},
            {"agent": "file", "action": "qwen_resume_optimization"},
            {"agent": "code", "action": "document_render"},
            {"agent": "file", "action": "register_docx_artifact"},
            {"agent": "supervisor", "action": "check_document_result"},
        ]
        return result

    def complete_resume_optimization_background(
        self,
        *,
        user: User,
        job_payload: dict[str, Any],
        phase_callback: ResumeJobPhaseCallback | None = None,
    ) -> dict[str, Any]:
        student = self._student_by_user(int(getattr(user, "id", 0) or 0))
        if not student:
            raise RuntimeError("当前账号未关联学生档案，无法继续后台简历优化。")

        attachment_id = self._to_int(job_payload.get("attachment_id"))
        attachment = self._attachment_by_id(student_id=int(student.id), attachment_id=attachment_id)
        if not attachment:
            raise RuntimeError("原简历附件已失效，请重新选择或上传简历。")
        file_path = self._resolve_attachment_abs_path(attachment)
        if not file_path:
            raise RuntimeError("原简历文件不存在，请重新上传简历后再试。")

        optimization_options = job_payload.get("optimization_options") if isinstance(job_payload.get("optimization_options"), dict) else {}
        message = str(job_payload.get("message") or "优化简历")
        timeout_seconds = int(getattr(self.settings, "FILE_AGENT_BACKGROUND_REQUEST_TIMEOUT_SECONDS", 180) or 180)
        max_retries = int(getattr(self.settings, "FILE_AGENT_BACKGROUND_MAX_RETRIES", 0) or 0)
        version = self._resume_version_for_attachment(student=student, attachment=attachment)

        self._emit_resume_job_phase(phase_callback, "extracting", "正在解析简历附件...")
        parsed = self._cached_parsed_resume(version)
        if parsed:
            extraction_report = parsed.get("file_extraction_report") if isinstance(parsed.get("file_extraction_report"), dict) else {}
            extract_steps = [{"tool": "code_file_extract_cache", "status": "done", "text": "done: reused cached local file extraction"}]
        else:
            extraction_report, extract_steps, extract_failure = self._extract_file_with_code_agent(
                user=user,
                student=student,
                attachment=attachment,
                file_path=file_path,
                message=message,
            )
            if extract_failure:
                raise RuntimeError(extract_failure)
            if self._is_extraction_report_low_quality(extraction_report, require_text=True):
                raise RuntimeError("CodeAgent local extraction returned empty resume text.")
            parsed = self._build_parsed_resume_from_extraction_report(attachment=attachment, extraction_report=extraction_report)
            self._store_parsed_resume_cache(version=version, parsed=parsed)

        self._emit_resume_job_phase(phase_callback, "optimizing", "正在按目标岗位优化简历...")
        cache_key = self._resume_optimization_cache_key(
            attachment_id=int(attachment.id),
            optimization_options=optimization_options,
        )
        cache_entry = self._cached_resume_optimization(version=version, cache_key=cache_key)
        cached_optimization = cache_entry.get("optimization_result") if isinstance(cache_entry.get("optimization_result"), dict) else None
        cached_artifact = self._cached_resume_export_artifact(student_id=int(student.id), cache_entry=cache_entry)
        result = self._complete_resume_optimization(
            user=user,
            student=student,
            attachment=attachment,
            parsed=parsed,
            extraction_report=extraction_report,
            extract_steps=extract_steps,
            message=message,
            optimization_options=optimization_options,
            timeout_seconds=timeout_seconds,
            max_retries=max_retries,
            optimization_result=cached_optimization,
            optimization_tool_steps=(
                [{"tool": "qwen_resume_optimization_cache", "status": "done", "text": "done: reused cached resume optimization"}]
                if cached_optimization
                else None
            ),
            cached_export_artifact=cached_artifact,
            phase_callback=phase_callback,
        )
        data = self._first_tool_output_data(result)
        optimization_result = data.get("resume_optimization_result") if isinstance(data.get("resume_optimization_result"), dict) else {}
        artifacts = data.get("artifacts") if isinstance(data.get("artifacts"), list) else []
        export_artifact = artifacts[0] if artifacts and isinstance(artifacts[0], dict) else {}
        if optimization_result:
            self._store_resume_optimization_cache(
                version=version,
                cache_key=cache_key,
                optimization_result=optimization_result,
                parsed=parsed,
                extraction_report=extraction_report,
                export_artifact=export_artifact,
            )
        return result

    def _build_resume_optimization_background_result(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        message: str,
        optimization_options: dict[str, Any],
    ) -> dict[str, Any]:
        job_id = f"resume_optimization_{uuid4().hex}"
        processing_message = "正在后台优化简历，我会在这条消息中自动更新最终优化稿。"
        context_patch = {
            "context_binding": {
                "attachment_id": attachment.id,
                "attachment": {
                    "id": attachment.id,
                    "file_name": attachment.file_name,
                    "file_type": attachment.file_type,
                },
                "resume_optimization_background_job": {
                    "id": job_id,
                    "status": "queued",
                    "phase": "queued",
                    "type": "resume_optimization",
                },
            }
        }
        background_job = {
            "id": job_id,
            "type": "resume_optimization",
            "status": "queued",
            "phase": "queued",
            "message": processing_message,
            "started_at": datetime.now().isoformat(timespec="seconds"),
            "payload": {
                "student_id": int(getattr(student, "id", 0) or 0),
                "attachment_id": int(getattr(attachment, "id", 0) or 0),
                "message": str(message or ""),
                "optimization_options": optimization_options,
                "wants_export": self._wants_resume_file_output(message),
            },
        }
        tool_output = self._tool_output(
            tool="optimize_resume",
            title="简历优化",
            summary=processing_message,
            data={
                "attachment_id": attachment.id,
                "background_job": {k: v for k, v in background_job.items() if k != "payload"},
            },
            next_actions=[],
            card_type="resume_card",
            context_patch=context_patch,
        )
        result = self._success(
            task_type="optimize_resume",
            reply=processing_message,
            tool_steps=[
                {"tool": "resume_optimization_background", "status": "queued", "text": "queued: resume optimization background job"},
            ],
            artifacts=[],
            context_patch=context_patch,
            tool_outputs=[tool_output],
        )
        result["background_job"] = background_job
        result["file_task"] = {
            "type": "optimize_resume",
            "status": "queued",
            "phase": "queued",
            "background": True,
            "background_job_id": job_id,
        }
        result["agent_flow_patch"] = [
            {"agent": "file", "action": "resolve_resume_attachment"},
            {"agent": "code", "action": "file_extract_request"},
            {"agent": "file", "action": "qwen_resume_optimization_background"},
            {"agent": "supervisor", "action": "return_processing_state"},
        ]
        return result

    @staticmethod
    def _emit_resume_job_phase(callback: ResumeJobPhaseCallback | None, phase: str, message: str) -> None:
        if callback:
            callback(phase, message)

    def _resume_version_for_attachment(self, *, student: Student, attachment: StudentAttachment) -> StudentResumeVersion:
        resume = (
            self.db.query(StudentResume)
            .options(joinedload(StudentResume.current_version))
            .filter(
                StudentResume.student_id == int(student.id),
                StudentResume.source_attachment_id == int(attachment.id),
                StudentResume.deleted.is_(False),
            )
            .order_by(StudentResume.is_default.desc(), StudentResume.id.desc())
            .first()
        )
        if not resume:
            default_exists = (
                self.db.query(StudentResume)
                .filter(
                    StudentResume.student_id == int(student.id),
                    StudentResume.deleted.is_(False),
                    StudentResume.is_default.is_(True),
                )
                .first()
                is not None
            )
            resume = StudentResume(
                student_id=int(student.id),
                title=Path(str(attachment.file_name or "resume")).stem or "resume",
                scene_type="uploaded_resume",
                is_default=not default_exists,
                status="active",
                source_attachment_id=int(attachment.id),
            )
            self.db.add(resume)
            self.db.flush()

        version = (
            self.db.query(StudentResumeVersion)
            .filter(
                StudentResumeVersion.resume_id == int(resume.id),
                StudentResumeVersion.attachment_id == int(attachment.id),
                StudentResumeVersion.deleted.is_(False),
            )
            .order_by(StudentResumeVersion.is_active.desc(), StudentResumeVersion.id.desc())
            .first()
        )
        if not version:
            version = StudentResumeVersion(
                resume_id=int(resume.id),
                version_no=len([item for item in list(resume.versions or []) if not getattr(item, "deleted", False)]) + 1,
                attachment_id=int(attachment.id),
                parsed_json={},
                optimized_json={},
                score_snapshot={},
                change_summary="resume processing cache",
                is_active=resume.current_version_id is None,
            )
            self.db.add(version)
            self.db.flush()
        if resume.current_version_id is None:
            resume.current_version_id = int(version.id)
            version.is_active = True
        return version

    def _cached_parsed_resume(self, version: StudentResumeVersion) -> dict[str, Any]:
        parsed = version.parsed_json if isinstance(version.parsed_json, dict) else {}
        extraction_report = parsed.get("file_extraction_report") if isinstance(parsed.get("file_extraction_report"), dict) else {}
        if extraction_report and not self._is_extraction_report_low_quality(extraction_report, require_text=True):
            return parsed
        return {}

    @staticmethod
    def _store_parsed_resume_cache(*, version: StudentResumeVersion, parsed: dict[str, Any]) -> None:
        version.parsed_json = parsed
        version.change_summary = "resume parsed cache"

    @staticmethod
    def _resume_optimization_cache_key(*, attachment_id: int, optimization_options: dict[str, Any]) -> str:
        payload = {
            "attachment_id": int(attachment_id),
            "target_role": str(optimization_options.get("target_role") or ""),
            "target_job_id": optimization_options.get("target_job_id"),
            "job_description": str(optimization_options.get("job_description") or ""),
            "render_template_version": "resume-blue-template-v2",
        }
        encoded = json.dumps(payload, sort_keys=True, ensure_ascii=False, default=str).encode("utf-8")
        return sha256(encoded).hexdigest()

    @staticmethod
    def _cached_resume_optimization(*, version: StudentResumeVersion, cache_key: str) -> dict[str, Any]:
        optimized = version.optimized_json if isinstance(version.optimized_json, dict) else {}
        cache = optimized.get("resume_optimization_cache") if isinstance(optimized.get("resume_optimization_cache"), dict) else {}
        entry = cache.get(cache_key) if isinstance(cache.get(cache_key), dict) else {}
        return entry

    @staticmethod
    def _store_resume_optimization_cache(
        *,
        version: StudentResumeVersion,
        cache_key: str,
        optimization_result: dict[str, Any],
        parsed: dict[str, Any],
        extraction_report: dict[str, Any],
        export_artifact: dict[str, Any],
    ) -> None:
        optimized = dict(version.optimized_json or {})
        cache = dict(optimized.get("resume_optimization_cache") or {})
        entry = {
            "optimization_result": optimization_result,
            "parsed_resume": parsed,
            "file_extraction_report": extraction_report,
            "updated_at": datetime.now().isoformat(timespec="seconds"),
        }
        if export_artifact:
            entry["export_artifact"] = export_artifact
        cache[cache_key] = entry
        optimized["resume_optimization_cache"] = cache
        version.optimized_json = optimized

    def _cached_resume_export_artifact(self, *, student_id: int, cache_entry: dict[str, Any]) -> dict[str, Any]:
        export_artifact = cache_entry.get("export_artifact") if isinstance(cache_entry.get("export_artifact"), dict) else {}
        artifact_id = self._to_int(export_artifact.get("id"))
        if not artifact_id:
            return {}
        attachment = self._attachment_by_id(student_id=student_id, attachment_id=artifact_id)
        if not attachment:
            return {}
        if not self._resolve_attachment_abs_path(attachment):
            return {}
        return {
            "id": attachment.id,
            "name": attachment.file_name,
            "type": attachment.file_type or "",
            "download_url": attachment.file_path,
            "mime_type": self._mime_type_by_suffix(Path(str(attachment.file_name or "")).suffix.lower()),
        }

    @staticmethod
    def _first_tool_output_data(result: dict[str, Any]) -> dict[str, Any]:
        for item in list(result.get("tool_outputs") or []):
            if isinstance(item, dict) and isinstance(item.get("data"), dict):
                return item["data"]
        return {}

    def _extract_file_with_code_agent(
        self,
        *,
        user: User | None,
        student: Student,
        attachment: StudentAttachment,
        file_path: Path,
        message: str,
    ) -> tuple[dict[str, Any], list[dict[str, Any]], str]:
        file_extract_request = {
            "source_path": str(file_path),
            "file_name": attachment.file_name,
            "file_type": attachment.file_type or Path(attachment.file_name).suffix.lstrip(".").lower(),
            "student_id": student.id,
            "source_attachment_id": attachment.id,
            "max_text_chars": 60000,
        }
        agent_user = user or SimpleNamespace(id=int(getattr(student, "user_id", 0) or 0), role=None, real_name=str(getattr(student, "name", "") or "student"))
        code_result = CodeAgent(self.db).execute(
            user=agent_user,
            message=self._build_code_agent_file_extract_prompt(file_extract_request=file_extract_request),
            selected_skill="code-agent",
            context_binding={"file_extract_request": file_extract_request},
            session_state={},
            client_state={},
        )
        report = self._extract_file_extraction_report(code_result)
        tool_steps = [
            {"tool": "code_file_extract", "status": "done", "text": "done: CodeAgent local file extraction"}
            if str(code_result.get("status") or "") == "success"
            else {"tool": "code_file_extract", "status": "failed", "text": "failed: CodeAgent local file extraction"}
        ]
        if str(code_result.get("status") or "") != "success":
            failure = str((code_result.get("code_task") or {}).get("failure_reason") or code_result.get("question") or "CodeAgent 本地文件解析失败。")
            return report, tool_steps, failure
        return report, tool_steps, ""

    @staticmethod
    def _build_code_agent_file_extract_prompt(*, file_extract_request: dict[str, Any]) -> str:
        return (
            "file_extract mode: run the controlled local Python extractor. "
            "Return a file_extraction_report with text, tables, images, charts, metadata, warnings, and source_file_type. "
            "Do not call an LLM to parse the document.\n"
            f"{json.dumps(file_extract_request, ensure_ascii=False)}"
        )

    @staticmethod
    def _extract_file_extraction_report(code_result: dict[str, Any]) -> dict[str, Any]:
        for item in list(code_result.get("tool_outputs") or []):
            if isinstance(item, dict):
                data = item.get("data") if isinstance(item.get("data"), dict) else {}
                report = data.get("file_extraction_report") if isinstance(data.get("file_extraction_report"), dict) else {}
                if report:
                    return report
        verification = code_result.get("verification")
        return verification if isinstance(verification, dict) else {}

    @staticmethod
    def _build_parsed_resume_from_extraction_report(
        *,
        attachment: StudentAttachment,
        extraction_report: dict[str, Any],
    ) -> dict[str, Any]:
        text = str(extraction_report.get("text") or "").strip()
        table_lines: list[str] = []
        for table in list(extraction_report.get("tables") or []):
            if not isinstance(table, dict):
                continue
            for row in list(table.get("rows") or []):
                if isinstance(row, list):
                    line = " | ".join(str(cell or "").strip() for cell in row if str(cell or "").strip())
                    if line:
                        table_lines.append(line)
        raw_text = "\n".join([part for part in [text, "\n".join(table_lines)] if part]).strip()
        return {
            "parser_success": True,
            "parser_engine": "code-agent-local-extract",
            "attachment_id": int(getattr(attachment, "id", 0) or 0),
            "attachment_name": str(getattr(attachment, "file_name", "") or ""),
            "attachment_type": str(getattr(attachment, "file_type", "") or ""),
            "source_file_type": str(extraction_report.get("source_file_type") or ""),
            "raw_text": raw_text,
            "text": text,
            "tables": list(extraction_report.get("tables") or []),
            "images": list(extraction_report.get("images") or []),
            "charts": list(extraction_report.get("charts") or []),
            "metadata": extraction_report.get("metadata") if isinstance(extraction_report.get("metadata"), dict) else {},
            "warnings": list(extraction_report.get("warnings") or []),
            "raw_text_length": len(raw_text),
            "file_extraction_report": extraction_report,
        }

    @staticmethod
    def _is_extraction_report_low_quality(extraction_report: dict[str, Any], *, require_text: bool) -> bool:
        if not isinstance(extraction_report, dict) or str(extraction_report.get("status") or "") != "success":
            return True
        text = str(extraction_report.get("text") or "").strip()
        table_text = json.dumps(extraction_report.get("tables") or [], ensure_ascii=False).strip("[]{}\" ")
        if text or table_text:
            return False
        if require_text:
            return True
        return not bool(extraction_report.get("images") or extraction_report.get("charts"))

    def _build_resume_optimization_result(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        extraction_report: dict[str, Any],
        message: str,
        optimization_options: dict[str, Any],
        timeout_seconds: int | None = None,
        max_retries: int | None = None,
    ) -> tuple[dict[str, Any], list[dict[str, Any]]]:
        system_prompt, user_prompt = self._build_resume_optimization_prompts(
            student=student,
            attachment=attachment,
            parsed=parsed,
            extraction_report=extraction_report,
            message=message,
            optimization_options=optimization_options,
        )
        resolved_timeout = int(timeout_seconds or getattr(self.settings, "FILE_AGENT_REQUEST_TIMEOUT_SECONDS", 60) or 60)
        raw_content = self._call_file_agent_json_model(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            timeout_seconds=resolved_timeout,
            max_retries=max_retries,
        )
        return self._load_resume_optimization_result(raw_content), [
            {"tool": "qwen_resume_optimization", "status": "done", "text": f"done: Qwen ResumeOptimizationResult within {resolved_timeout}s"}
        ]

    def _build_resume_optimization_prompts(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        extraction_report: dict[str, Any],
        message: str,
        optimization_options: dict[str, Any],
    ) -> tuple[str, str]:
        system_prompt = (
            "You are the File Agent resume optimizer. Return ONLY one strict JSON object, no markdown. "
            "The JSON is ResumeOptimizationResult with fields: target_role, reply_summary, "
            "optimized_resume_document, revision_suggestions, job_match_advice, missing_evidence, and career_planning_advice. "
            "Do not produce DocumentBuildSpec. Do not invent projects, awards, certificates, grades, employers, dates, or metrics. "
            "If evidence is missing, put it in missing_evidence or evidence_boundary. "
            "CRITICAL: optimized_resume_document must contain ONLY pure resume content ready for direct submission. "
            "Do NOT include any explanatory, advisory, or instructional text in any field value. "
            "Forbidden patterns in field values: \u5efa\u8bae, \u91cd\u70b9, \u6ce8\u610f, \u9700\u8981, \u5e94\u8be5, suggest, recommend, note, important, should. "
            "Each rewrite field must be formal professional text suitable for a real resume, not analysis or coaching notes."
        )
        payload = {
            "student_profile": self._student_prompt_payload(student),
            "source_attachment": {
                "id": attachment.id,
                "file_name": attachment.file_name,
                "file_type": attachment.file_type,
            },
            "user_request": message,
            "optimization_options": optimization_options,
            "local_file_extraction_report": self._compact_prompt_payload(extraction_report),
            "parsed_resume_text": str(parsed.get("raw_text") or "")[:12000],
            "required_output": "ResumeOptimizationResult JSON only; no DOCX render spec.",
        }
        user_prompt = (
            "Optimize the resume based only on the local extraction report below. "
            "Keep the result compact and suitable for either text reply or deterministic DOCX rendering. "
            "In optimized_resume_document: write only submission-ready resume content. "
            "No coaching notes, no suggestions, no meta-commentary in any field value.\n"
            f"{json.dumps(payload, ensure_ascii=False)}"
        )
        return system_prompt, user_prompt

    def _load_resume_optimization_result(self, raw_content: str) -> dict[str, Any]:
        result = json.loads(raw_content)
        if not isinstance(result, dict):
            raise ValueError("ResumeOptimizationResult is not a JSON object")
        optimized = result.get("optimized_resume_document") or result.get("optimized_resume_text") or result.get("optimized_resume")
        if not optimized:
            raise ValueError("ResumeOptimizationResult missing optimized_resume_document")
        if "optimized_resume_document" not in result:
            result["optimized_resume_document"] = optimized
        result["missing_evidence"] = self._normalize_string_list(result.get("missing_evidence"))
        suggestions = self._normalize_resume_revision_suggestions(
            self._resume_revision_suggestions_source(result)
        )
        if not suggestions:
            suggestions = self._fallback_resume_revision_suggestions(result["missing_evidence"])
        result["revision_suggestions"] = suggestions
        result["job_match_advice"] = self._normalize_advice_block(result.get("job_match_advice"), default="结合目标岗位要求补强项目、技能和成果证据。")
        result["career_planning_advice"] = self._normalize_advice_block(result.get("career_planning_advice"), default="按目标岗位能力模型补齐项目深度、实习/竞赛证据和面试表达。")
        result["reply_summary"] = str(result.get("reply_summary") or "已基于本地解析内容完成简历优化建议。").strip()
        result["target_role"] = str(result.get("target_role") or "").strip()
        return result

    @staticmethod
    def _normalize_advice_block(value: Any, *, default: str) -> Any:
        if isinstance(value, (dict, list)) and value:
            return value
        text = str(value or "").strip()
        return text or default

    @staticmethod
    def _normalize_string_list(value: Any) -> list[str]:
        if isinstance(value, list):
            return [str(item or "").strip() for item in value if str(item or "").strip()]
        text = str(value or "").strip()
        return [text] if text else []

    @staticmethod
    def _resume_revision_suggestions_source(result: dict[str, Any]) -> Any:
        for key in (
            "revision_suggestions",
            "suggestions",
            "optimization_suggestions",
            "improvement_suggestions",
            "optimization_notes",
        ):
            value = result.get(key)
            if value not in (None, "", [], {}):
                return value
        return result.get("revision_suggestions")

    @staticmethod
    def _normalize_resume_revision_suggestions(value: Any) -> list[dict[str, str]]:
        if isinstance(value, list):
            rows = value
        elif isinstance(value, dict):
            direct_keys = {
                "field",
                "title",
                "suggestion",
                "text",
                "content",
                "advice",
                "note",
                "notes",
                "detail",
                "description",
                "evidence_boundary",
                "missing_evidence",
                "enterprise_reason",
                "career_reason",
            }
            if direct_keys.intersection(value.keys()):
                rows = [value]
            else:
                rows = []
                for key, item in value.items():
                    if isinstance(item, list):
                        rows.extend(item)
                    elif isinstance(item, dict):
                        row = dict(item)
                        row.setdefault("field", str(key or "").strip())
                        rows.append(row)
                    else:
                        rows.append({"field": str(key or "").strip(), "suggestion": str(item or "").strip()})
        elif isinstance(value, str):
            rows = [value]
        else:
            rows = []
        suggestions: list[dict[str, str]] = []
        for index, item in enumerate(rows, start=1):
            if isinstance(item, dict):
                field = str(item.get("field") or item.get("title") or item.get("area") or item.get("section") or item.get("category") or f"建议 {index}").strip()
                suggestion = str(
                    item.get("suggestion")
                    or item.get("text")
                    or item.get("content")
                    or item.get("advice")
                    or item.get("note")
                    or item.get("notes")
                    or item.get("detail")
                    or item.get("description")
                    or ""
                ).strip()
                if not suggestion:
                    scalar_values = [
                        str(nested or "").strip()
                        for key, nested in item.items()
                        if key
                        not in {
                            "field",
                            "title",
                            "area",
                            "section",
                            "category",
                            "enterprise_reason",
                            "career_reason",
                            "evidence_boundary",
                            "missing_evidence",
                        }
                        and not isinstance(nested, (dict, list))
                        and str(nested or "").strip()
                    ]
                    suggestion = "；".join(scalar_values[:3])
                evidence_boundary = str(item.get("evidence_boundary") or item.get("missing_evidence") or "仅基于已解析简历内容，新增事实需学生确认。").strip()
                enterprise_reason = str(item.get("enterprise_reason") or item.get("reason") or "提升与目标岗位要求的匹配度。").strip()
                career_reason = str(item.get("career_reason") or "帮助后续职业规划与面试表达更聚焦。").strip()
            else:
                field = f"建议 {index}"
                suggestion = str(item or "").strip()
                enterprise_reason = "提升与目标岗位要求的匹配度。"
                career_reason = "帮助后续职业规划与面试表达更聚焦。"
                evidence_boundary = "仅基于已解析简历内容，新增事实需学生确认。"
            if suggestion:
                suggestions.append(
                    {
                        "field": field or f"建议 {index}",
                        "suggestion": suggestion,
                        "enterprise_reason": enterprise_reason,
                        "career_reason": career_reason,
                        "evidence_boundary": evidence_boundary,
                    }
                )
        return suggestions

    @staticmethod
    def _fallback_resume_revision_suggestions(missing_evidence: list[str]) -> list[dict[str, str]]:
        if missing_evidence:
            evidence_text = "；".join(missing_evidence[:3])
            return [
                {
                    "field": "证据补充",
                    "suggestion": f"补充或确认以下可验证材料：{evidence_text}",
                    "enterprise_reason": "确保投递内容有事实依据，提升岗位匹配表达的可信度。",
                    "career_reason": "帮助后续面试复盘和职业规划材料保持一致。",
                    "evidence_boundary": "仅基于模型返回的缺失证据清单生成提醒，不新增简历事实。",
                }
            ]
        return [
            {
                "field": "整体校对",
                "suggestion": "核对优化稿中的经历、项目、证书、时间线与量化成果是否真实准确，确认后再投递。",
                "enterprise_reason": "保证简历内容真实可信，降低投递和背调风险。",
                "career_reason": "帮助后续面试表达与职业规划材料保持一致。",
                "evidence_boundary": "仅提醒核对既有优化稿，不新增简历事实。",
            }
        ]

    def _format_resume_optimization_reply(self, optimization_result: dict[str, Any], *, fallback_summary: str) -> str:
        lines = [str(optimization_result.get("reply_summary") or fallback_summary).strip()]
        target_role = str(optimization_result.get("target_role") or "").strip()
        if target_role:
            lines.append(f"\n目标岗位：{target_role}")
        optimized_resume = self._format_optimized_resume_document(optimization_result.get("optimized_resume_document"))
        if optimized_resume:
            lines.append("\n优化后简历：")
            lines.append(optimized_resume)
        suggestions = list(optimization_result.get("revision_suggestions") or [])
        if suggestions:
            lines.append("\n修改建议：")
            for item in suggestions[:6]:
                field = str(item.get("field") or "建议").strip()
                suggestion = str(item.get("suggestion") or "").strip()
                if suggestion:
                    lines.append(f"- {field}：{suggestion}")
        missing = list(optimization_result.get("missing_evidence") or [])
        if missing:
            lines.append("\n需要你补充或确认的证据：")
            for item in missing[:5]:
                lines.append(f"- {item}")
        attachment_id = optimization_result.get("attachment_id")
        if attachment_id:
            lines.append("\n## 📄 导出文件")
            lines.append("| 格式 | 下载链接 |")
            lines.append("|------|----------|")
            lines.append(f"| Word (.docx) | [点击下载](/api/students/me/resume/export/word/{attachment_id}) |")
            lines.append(f"| PDF (.pdf) | [点击下载](/api/students/me/resume/export/pdf/{attachment_id}) |")
        lines.append('\n如需生成文件，请继续发送"优化并导出 Word"。')
        return "\n".join(lines)

    @classmethod
    def _format_optimized_resume_document(cls, value: Any) -> str:
        lines: list[str] = []

        def append(item: Any, *, indent: int = 0, label: str = "") -> None:
            prefix = "  " * indent
            if isinstance(item, dict):
                if label:
                    lines.append(f"{prefix}{label}：")
                for key, nested in item.items():
                    title = str(key or "").strip()
                    if isinstance(nested, (dict, list)):
                        append(nested, indent=indent + (1 if label else 0), label=title)
                    else:
                        text = str(nested or "").strip()
                        if title and text:
                            lines.append(f"{prefix}{'  ' if label else ''}{title}：{text}")
                        elif text:
                            lines.append(f"{prefix}{'  ' if label else ''}{text}")
                return

            if isinstance(item, list):
                if label:
                    lines.append(f"{prefix}{label}：")
                item_prefix = "  " * (indent + (1 if label else 0))
                for nested in item:
                    if isinstance(nested, (dict, list)):
                        append(nested, indent=indent + (1 if label else 0), label="")
                    else:
                        text = str(nested or "").strip()
                        if text:
                            lines.append(f"{item_prefix}- {text}")
                return

            text = str(item or "").strip()
            if not text:
                return
            lines.append(f"{prefix}{label}：{text}" if label else f"{prefix}{text}")

        append(value)
        return "\n".join(line for line in lines if str(line or "").strip()).strip()

    @staticmethod
    def _wants_resume_file_output(message: str) -> bool:
        compact = re.sub(r"\s+", "", str(message or "").lower())
        action_tokens = ("导出", "下载", "生成", "创建", "保存", "export", "download", "generate", "create", "save")
        file_tokens = ("word", "docx", "doc", "pdf", "文件", "文档")
        return any(token in compact for token in action_tokens) and any(token in compact for token in file_tokens)

    def _build_document_spec_from_resume_optimization(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        optimization_result: dict[str, Any],
        optimization_options: dict[str, Any],
    ) -> dict[str, Any]:
        target_role = (
            str(optimization_result.get("target_role") or "").strip()
            or str(optimization_options.get("target_role") or "").strip()
            or str(getattr(student, "target_industry", "") or "").strip()
            or "目标岗位"
        )
        suggestions = self._normalize_resume_revision_suggestions(optimization_result.get("revision_suggestions"))
        optimized_document = optimization_result.get("optimized_resume_document")
        if not isinstance(optimized_document, (dict, list)):
            optimized_document = {"正文": str(optimized_document or parsed.get("raw_text") or "").strip()}
        spec = {
            "document_type": "optimized_resume",
            "target_role": target_role,
            "enterprise_alignment": optimization_result.get("job_match_advice")
            or {"summary": "围绕目标企业与岗位要求，突出可验证项目、技术栈和成果证据。"},
            "career_planning": optimization_result.get("career_planning_advice")
            or {"summary": "围绕目标岗位补齐项目深度、实习实践、证书与面试表达。"},
            "revision_suggestions": suggestions,
            "optimized_resume_document": optimized_document,
            "document_style": {
                "colors": ["deep blue", "teal green", "warm orange"],
                "density": "compact",
                "source": "ResumeOptimizationResult",
                "fonts": {
                    "chinese_default": "宋体 (SimSun)",
                    "western_default": "Times New Roman",
                    "heading_font": "黑体 (SimHei)",
                    "code_font": "Consolas"
                },
                "headings": {
                    "h1": {"size": 16, "bold": True, "font": "SimHei", "spacing_before": 12, "spacing_after": 6},
                    "h2": {"size": 14, "bold": True, "font": "SimHei", "spacing_before": 10, "spacing_after": 4},
                    "h3": {"size": 12, "bold": True, "font": "KaiTi", "spacing_before": 8, "spacing_after": 2},
                    "h4": {"size": 12, "bold": True, "font": "SimSun", "spacing_before": 6, "spacing_after": 0},
                    "h5": {"size": 10.5, "bold": False, "font": "SimSun", "spacing_before": 4, "spacing_after": 0}
                },
                "table_style": {
                    "border_width": 1,
                    "border_color": "#000000",
                    "header_bg": "#1e40af",
                    "header_text_color": "#ffffff",
                    "header_font_size": 11,
                    "row_alternate_bg": "#f8f9fa",
                    "cell_padding_v": 6,
                    "cell_padding_h": 8
                },
                "body_text": {
                    "font_size": 12,
                    "line_spacing": 1.5,
                    "first_line_indent_chars": 2
                }
            },
            "sections": [
                {"id": "enterprise_alignment", "title": "企业匹配分析", "content": optimization_result.get("job_match_advice") or ""},
                {"id": "career_planning_path", "title": "职业规划路径", "content": optimization_result.get("career_planning_advice") or ""},
                {"id": "revision_suggestions_table", "title": "修改建议表", "content": suggestions},
                {"id": "optimized_resume", "title": "优化后简历正文", "content": optimized_document},
            ],
            "supervisor_checklist": [
                "核对所有经历、项目、证书和时间线是否真实。",
                "补充 missing_evidence 中列出的可验证材料后再投递。",
                f"确认导出的文件来自附件：{attachment.file_name}",
            ],
        }
        self._validate_document_build_spec(spec)
        return spec

    def _build_document_build_spec(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        message: str,
        optimization_options: dict[str, Any],
    ) -> tuple[dict[str, Any], list[dict[str, Any]]]:
        system_prompt, user_prompt = self._build_document_build_spec_prompts(
            student=student,
            attachment=attachment,
            parsed=parsed,
            message=message,
            optimization_options=optimization_options,
        )
        timeout_seconds = int(getattr(self.settings, "FILE_AGENT_DOCUMENT_SPEC_TIMEOUT_SECONDS", 90) or 90)
        retry_count = max(0, int(getattr(self.settings, "FILE_AGENT_DOCUMENT_SPEC_RETRY_COUNT", 1) or 0))
        tool_steps: list[dict[str, Any]] = []
        last_timeout: TimeoutError | None = None

        try:
            raw_content = self._call_file_agent_json_model(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                timeout_seconds=timeout_seconds,
            )
            return self._load_document_build_spec(raw_content), [
                {"tool": "qwen_document_spec", "status": "done", "text": f"done: Qwen DocumentBuildSpec within {timeout_seconds}s"}
            ]
        except TimeoutError as exc:
            last_timeout = exc
            tool_steps.append(
                {"tool": "qwen_document_spec", "status": "failed", "text": f"failed: Qwen DocumentBuildSpec timed out after {timeout_seconds}s"}
            )

        for retry_index in range(1, retry_count + 1):
            compact_system_prompt, compact_user_prompt = self._build_compact_document_build_spec_prompts(
                student=student,
                attachment=attachment,
                parsed=parsed,
                message=message,
                optimization_options=optimization_options,
            )
            try:
                raw_content = self._call_file_agent_json_model(
                    system_prompt=compact_system_prompt,
                    user_prompt=compact_user_prompt,
                    timeout_seconds=timeout_seconds,
                )
                tool_steps.append(
                    {
                        "tool": "qwen_document_spec_retry",
                        "status": "done",
                        "text": f"done: compact DocumentBuildSpec retry {retry_index}",
                    }
                )
                return self._load_document_build_spec(raw_content), tool_steps
            except TimeoutError as exc:
                last_timeout = exc
                tool_steps.append(
                    {
                        "tool": "qwen_document_spec_retry",
                        "status": "failed",
                        "text": f"failed: compact DocumentBuildSpec retry {retry_index} timed out after {timeout_seconds}s",
                    }
                )

        raise DocumentBuildSpecTimeoutError(
            (
                f"文档构建规范生成超时。AI模型处理时间过长，可能原因：\n"
                f"1. 简历内容较为复杂或篇幅较长\n"
                f"2. 当前API服务负载较高\n"
                f"3. 网络连接不稳定\n\n"
                f"建议操作：\n"
                f"- 检查网络连接\n"
                f"- 稍后重新尝试（建议等待30秒以上）\n"
                f"- 如持续出现问题，可联系管理员调整超时设置"
            ),
            tool_steps=tool_steps,
        )

    def _load_document_build_spec(self, raw_content: str) -> dict[str, Any]:
        spec = json.loads(raw_content)
        if not isinstance(spec, dict):
            raise ValueError("Qwen response is not a JSON object")
        self._validate_document_build_spec(spec)
        return spec

    def _build_document_build_spec_prompts(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        message: str,
        optimization_options: dict[str, Any],
    ) -> tuple[str, str]:
        system_prompt = (
            "You are the File Agent for a college career-planning platform. "
            "Return ONLY one strict JSON object named by its fields, with no markdown and no wrapper key. "
            "The object must be a DocumentBuildSpec with these top-level fields: "
            f"{', '.join(self.DOCUMENT_BUILD_SPEC_REQUIRED_FIELDS)}. "
            "The content must be strongly related to enterprise hiring fit and career planning. "
            "Every item in revision_suggestions must include field, suggestion, enterprise_reason, career_reason, and evidence_boundary. "
            "Do not invent experiences, companies, certificates, time lines, grades, awards, or quantified outcomes. "
            "If evidence is missing, write a revision suggestion that asks the student to verify or add evidence instead of fabricating it. "
            "Set document_style colors to deep blue, teal green, and warm orange. "
            "sections must include cover, summary_card, enterprise_alignment, career_planning_path, revision_suggestions_table, and action_checklist, "
            "and must expose the Chinese section titles 企业匹配分析, 职业规划路径, 修改建议表.\n\n"
            "--- Document Format Specifications ---\n\n"
            "【字体规范】\n"
            "- 中文默认字体：宋体 (SimSun)\n"
            "- 西文/数字默认字体：Times New Roman\n"
            "- 标题字体：黑体(SimHei)用于H1-H2，楷体(KaiTi)用于H3，宋体用于H4-H5\n"
            "- 代码块字体：Consolas，10pt\n\n"
            "【标题层级规范（5级）】\n"
            "- H1: 16pt(小二号), 黑体, 加粗, 段前12pt段后6pt\n"
            "- H2: 14pt(四号), 黑体, 加粗, 段前10pt段后4pt\n"
            "- H3: 12pt(小四号), 楷体/黑体, 加粗, 段前8pt段后2pt\n"
            "- H4: 12pt(小四号), 宋体, 加粗, 段前6pt段后0pt\n"
            "- H5: 10.5pt(五号), 宋体, 常规, 段前4pt段后0pt\n\n"
            "【表格样式规范】\n"
            "- 边框：1pt实线黑色(#000000)\n"
            "- 表头：深蓝背景(#1e40af)，白色文字，黑体11pt，加粗\n"
            "- 数据行：斑马纹效果（奇数行白色#ffffff，偶数行浅灰#f8f9fa）\n"
            "- 单元格内边距：上下6pt，左右8pt\n"
            "- 对齐方式：水平垂直居中\n\n"
            "【正文字体格式】\n"
            "- 字号：12pt(小四号)\n"
            "- 行距：1.5倍\n"
            "- 首行缩进：2字符\n"
            "- 中文字体：宋体 + 西文字体Times New Roman混合\n\n"
            "--- Compact Format Summary ---\n"
            "Document formatting requirements (compact):\n"
            "- Fonts: Chinese=宋体(SimSun), Western=Times New Roman, Headings H1-H2=黑体, H3=楷体, H4-H5=宋体\n"
            "- Headings: H1(16pt/Bold/SimHei), H2(14pt/Bold/SimHei), H3(12pt/Bold/KaiTi), H4(12pt/Bold/SimSun), H5(10.5pt/Normal/SimSun)\n"
            "- Table: border=1pt black, header_bg=#1e40af/white text, zebra rows(white/#f8f9fa), padding=6v8h\n"
            "- Body: 12pt, line-spacing=1.5x, indent=2chars"
        )
        payload = {
            "student_profile": self._student_prompt_payload(student),
            "source_attachment": {
                "id": attachment.id,
                "file_name": attachment.file_name,
                "file_type": attachment.file_type,
            },
            "user_request": message,
            "optimization_options": optimization_options,
            "local_parse_result": parsed,
            "required_output_format": "docx",
        }
        user_prompt = (
            "Use the local parse result below to build the DocumentBuildSpec. "
            "Keep the optimized resume faithful to the parsed resume and make all changes traceable to enterprise fit or career planning.\n"
            f"{json.dumps(payload, ensure_ascii=False)}"
        )
        return system_prompt, user_prompt

    def _build_compact_document_build_spec_prompts(
        self,
        *,
        student: Student,
        attachment: StudentAttachment,
        parsed: dict[str, Any],
        message: str,
        optimization_options: dict[str, Any],
    ) -> tuple[str, str]:
        system_prompt = (
            "Return ONLY one strict JSON DocumentBuildSpec object, no markdown. "
            f"Required top-level fields: {', '.join(self.DOCUMENT_BUILD_SPEC_REQUIRED_FIELDS)}. "
            "Every revision_suggestions item must include field, suggestion, enterprise_reason, career_reason, and evidence_boundary. "
            "Do not invent facts; ask for verification where evidence is missing. "
            "sections must include Chinese titles 企业匹配分析, 职业规划路径, 修改建议表. "
            "document_style colors must be deep blue, teal green, and warm orange."
        )
        payload = {
            "student_profile": self._compact_prompt_payload(self._student_prompt_payload(student)),
            "source_attachment": {
                "id": attachment.id,
                "file_name": attachment.file_name,
                "file_type": attachment.file_type,
            },
            "user_request": str(message or "")[:800],
            "optimization_options": self._compact_prompt_payload(optimization_options),
            "local_parse_result": self._compact_prompt_payload(parsed),
            "required_output_format": "docx",
            "retry_instruction": "This is a retry after a read timeout. Produce a concise but complete valid DocumentBuildSpec.",
        }
        user_prompt = (
            "Build the DocumentBuildSpec from this compact payload. Keep all changes traceable to enterprise fit or career planning.\n"
            f"{json.dumps(payload, ensure_ascii=False)}"
        )
        return system_prompt, user_prompt

    def _compact_prompt_payload(self, value: Any, *, depth: int = 0) -> Any:
        if depth >= 4:
            return str(value)[:500]
        if isinstance(value, dict):
            return {str(key): self._compact_prompt_payload(item, depth=depth + 1) for key, item in list(value.items())[:20]}
        if isinstance(value, list):
            return [self._compact_prompt_payload(item, depth=depth + 1) for item in value[:12]]
        if isinstance(value, str):
            return value[:1600]
        return value

    def _student_prompt_payload(self, student: Student) -> dict[str, Any]:
        return {
            "id": student.id,
            "name": student.name,
            "grade": student.grade,
            "major": student.major,
            "college": student.college,
            "target_industry": student.target_industry,
            "target_city": student.target_city,
            "education_experience": student.education_experience,
            "bio": student.bio,
            "skills": [
                {"name": item.name, "level": item.level, "category": item.category, "description": item.description}
                for item in list(student.skills or [])[:30]
            ],
            "certificates": [
                {"name": item.name, "issuer": item.issuer, "issued_date": item.issued_date, "description": item.description}
                for item in list(student.certificates or [])[:30]
            ],
            "projects": [
                {"name": item.name, "role": item.role, "description": item.description, "outcome": item.outcome}
                for item in list(student.projects or [])[:20]
            ],
            "internships": [
                {"company": item.company, "position": item.position, "description": item.description}
                for item in list(student.internships or [])[:20]
            ],
        }

    def _call_file_agent_json_model(self, *, system_prompt: str, user_prompt: str, timeout_seconds: int | None = None, max_retries: int | None = None) -> str:
        api_key = str(self.llm_profile.api_key or "").strip()
        base_url = str(self.llm_profile.base_url or "").strip().rstrip("/")
        model_name = str(self.llm_profile.module_name or "").strip()
        if not api_key:
            raise RuntimeError("FILE_AGENT_API_KEY is not configured")
        if not base_url:
            raise RuntimeError("FILE_AGENT_BASE_URL is not configured")
        if not model_name:
            raise RuntimeError("FILE_AGENT_MODULE_NAME is not configured")

        endpoint = base_url if base_url.endswith("/chat/completions") else f"{base_url}/chat/completions"
        payload = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.2,
            "response_format": {"type": "json_object"},
        }
        request_data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        if max_retries is None:
            max_retries = int(getattr(self.settings, "FILE_AGENT_MAX_RETRIES", 0) or 0)
        else:
            max_retries = int(max_retries)
        max_retries = max(0, max_retries)
        transient_network_max_retries = max(
            0,
            int(getattr(self.settings, "FILE_AGENT_TRANSIENT_NETWORK_MAX_RETRIES", 1) or 0),
        )
        attempt_retry_limit = max(max_retries, transient_network_max_retries)
        base_timeout = int(timeout_seconds or getattr(self.settings, 'FILE_AGENT_REQUEST_TIMEOUT_SECONDS', None) or 60)
        last_exception = None

        logger = logging.getLogger(__name__)
        start_time = time.time()
        prompt_length = len(system_prompt) + len(user_prompt)
        logger.info(
            f"FileAgent LLM API call started | "
            f"timeout={base_timeout}s | "
            f"max_retries={max_retries} | "
            f"transient_network_max_retries={transient_network_max_retries} | "
            f"system_prompt_length={len(system_prompt)} chars | "
            f"user_prompt_length={len(user_prompt)} chars"
        )

        for attempt in range(attempt_retry_limit + 1):
            try:
                request = Request(
                    endpoint,
                    data=request_data,
                    headers=headers,
                    method="POST",
                )
                with urlopen(request, timeout=base_timeout) as response:
                    response_data = json.loads(response.read().decode("utf-8"))
            except (TimeoutError, socket.timeout) as exc:
                last_exception = exc
                if attempt < max_retries:
                    wait_time = 2 ** (attempt + 1)
                    logger.warning(
                        f"FileAgent LLM API call timed out (attempt {attempt+1}/{max_retries+1}, "
                        f"timeout={base_timeout}s). Retrying in {wait_time}s..."
                    )
                    time.sleep(wait_time)
                else:
                    total_elapsed = time.time() - start_time
                    logger.error(
                        f"FileAgent LLM API call failed after all retries | "
                        f"total_attempts={max_retries+1} | "
                        f"total_elapsed_time={total_elapsed:.2f}s | "
                        f"timeout_setting={base_timeout}s | "
                        f"last_error={str(last_exception)}"
                    )
                    raise TimeoutError(
                        f"LLM API call failed after {max_retries} retries. "
                        f"Last error: {str(exc)}"
                    ) from exc
            except URLError as exc:
                if self._is_timeout_error(exc):
                    last_exception = exc
                    if attempt < max_retries:
                        wait_time = 2 ** (attempt + 1)
                        logger.warning(
                            f"FileAgent LLM API call timed out (attempt {attempt+1}/{max_retries+1}, "
                            f"timeout={base_timeout}s). Retrying in {wait_time}s..."
                        )
                        time.sleep(wait_time)
                        continue
                    else:
                        total_elapsed = time.time() - start_time
                        logger.error(
                            f"FileAgent LLM API call failed after all retries | "
                            f"total_attempts={max_retries+1} | "
                            f"total_elapsed_time={total_elapsed:.2f}s | "
                            f"timeout_setting={base_timeout}s | "
                            f"last_error={str(last_exception)}"
                        )
                        raise TimeoutError(
                            f"LLM API call failed after {max_retries} retries. "
                            f"Last error: {str(exc)}"
                        ) from exc
                if self._is_transient_network_error(exc):
                    last_exception = exc
                    retry_limit = max(max_retries, transient_network_max_retries)
                    if attempt < retry_limit:
                        wait_time = 2 ** (attempt + 1)
                        logger.warning(
                            f"FileAgent LLM API call hit a transient network error "
                            f"(attempt {attempt+1}/{retry_limit+1}, timeout={base_timeout}s): "
                            f"{self._network_error_summary(exc)}. Retrying in {wait_time}s..."
                        )
                        time.sleep(wait_time)
                        continue
                    else:
                        total_elapsed = time.time() - start_time
                        logger.error(
                            f"FileAgent LLM API call failed after transient network retries | "
                            f"total_attempts={retry_limit+1} | "
                            f"total_elapsed_time={total_elapsed:.2f}s | "
                            f"timeout_setting={base_timeout}s | "
                            f"last_error={str(last_exception)}"
                        )
                        raise TimeoutError(
                            f"LLM API transient network error after {retry_limit} retries: "
                            f"{self._network_error_summary(exc)}"
                        ) from exc
                raise

        content = response_data["choices"][0]["message"]["content"]
        if not isinstance(content, str):
            raise ValueError("Qwen JSON content is not text")

        elapsed_time = time.time() - start_time
        logger.info(
            f"FileAgent LLM API call succeeded | "
            f"attempt={attempt+1} | "
            f"elapsed_time={elapsed_time:.2f}s | "
            f"timeout={base_timeout}s | "
            f"response_length={len(content)} chars"
        )

        return content

    @staticmethod
    def _is_timeout_error(exc: BaseException) -> bool:
        reason = getattr(exc, "reason", None)
        if isinstance(reason, (TimeoutError, socket.timeout)):
            return True
        text = str(exc).lower()
        return "timed out" in text or "timeout" in text

    @staticmethod
    def _is_transient_network_error(exc: BaseException) -> bool:
        reason = getattr(exc, "reason", None)
        text = str(reason or exc).lower()
        if hasattr(exc, "code"):
            return False
        if isinstance(reason, (ConnectionResetError, ConnectionAbortedError, BrokenPipeError, ssl.SSLError)):
            return True
        transient_markers = (
            "unexpected_eof_while_reading",
            "eof occurred in violation of protocol",
            "remote end closed connection",
            "connection reset",
            "connection aborted",
            "broken pipe",
            "temporarily unavailable",
        )
        return any(marker in text for marker in transient_markers)

    @staticmethod
    def _network_error_summary(exc: BaseException) -> str:
        reason = getattr(exc, "reason", None)
        text = str(reason or exc).strip() or str(exc)
        lowered = text.lower()
        if "unexpected_eof_while_reading" in lowered or "eof occurred in violation of protocol" in lowered:
            return "HTTPS 连接被远端提前关闭（SSL EOF）"
        if "connection reset" in lowered:
            return "HTTPS 连接被重置"
        if "connection aborted" in lowered:
            return "HTTPS 连接被中断"
        return f"HTTPS 网络连接异常：{text}"[:300]

    def _validate_document_build_spec(self, spec: dict[str, Any]) -> None:
        missing = [field for field in self.DOCUMENT_BUILD_SPEC_REQUIRED_FIELDS if field not in spec or spec[field] in ("", None, [], {})]
        if missing:
            raise ValueError(f"DocumentBuildSpec missing required fields: {', '.join(missing)}")
        revision_suggestions = spec["revision_suggestions"]
        if not isinstance(revision_suggestions, list) or not revision_suggestions:
            raise ValueError("revision_suggestions must be a non-empty list")
        for index, item in enumerate(revision_suggestions, start=1):
            if not isinstance(item, dict):
                raise ValueError(f"revision_suggestions[{index}] is not an object")
            if not str(item.get("enterprise_reason") or "").strip():
                raise ValueError(f"revision_suggestions[{index}] missing enterprise_reason")
            if not str(item.get("career_reason") or "").strip():
                raise ValueError(f"revision_suggestions[{index}] missing career_reason")
            if not str(item.get("field") or "").strip():
                raise ValueError(f"revision_suggestions[{index}] missing field")
            if not str(item.get("suggestion") or "").strip():
                raise ValueError(f"revision_suggestions[{index}] missing suggestion")
            if not str(item.get("evidence_boundary") or "").strip():
                raise ValueError(f"revision_suggestions[{index}] missing evidence_boundary")
        section_text = json.dumps(spec["sections"], ensure_ascii=False)
        for title in ("企业匹配分析", "职业规划", "修改建议"):
            if title not in section_text:
                raise ValueError(f"sections missing required chapter title: {title}")

    @staticmethod
    def _build_code_agent_document_render_prompt(*, document_render_request: dict[str, Any]) -> str:
        return (
            "document_render mode: use python-docx to generate a submission-ready resume DOCX. "
            "Use only optimized_resume_document as final resume content; do not render summary cards, "
            "enterprise/career analysis, revision tables, action checklists, JSON, or agent flow text. "
            "Follow a clean blue-white resume template: top blue ribbon, photo placeholder, large blue name, "
            "two-column contact grid, compact section labels, and dense professional resume body text. "
            "Run render tests that verify the DOCX exists, is larger than 5KB, contains resume section titles, "
            "and excludes report-only chapter titles. "
            "Do not generate a substitute or mock document.\n\n"
            "=== python-docx Rendering Parameter Guidelines ===\n\n"
            "**1. Font Settings (Chinese & Western text):**\n"
            "```python\n"
            "# Font settings for Chinese and Western text\n"
            "run.font.name = 'Times New Roman'\n"
            "run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')\n"
            "style.font.size = Pt(12)  # 小四号\n"
            "```\n\n"
            "**2. Heading Styles (H1-H5):**\n"
            "```python\n"
            "# Heading styles\n"
            "for level in range(1, 6):\n"
            "    style = doc.styles[f'Heading {level}']\n"
            "    font = style.font\n"
            "    if level == 1:\n"
            "        font.size = Pt(16); font.bold = True; font.name = 'SimHei'\n"
            "        font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')\n"
            "    elif level == 2:\n"
            "        font.size = Pt(14); font.bold = True; font.name = 'SimHei'\n"
            "        font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')\n"
            "    # ... H3-H5 similar\n"
            "```\n\n"
            "**3. Table Styling:**\n"
            "```python\n"
            "# Table styling\n"
            "table.style = 'Table Grid'\n"
            "for row_idx, row in enumerate(table.rows):\n"
            "    for cell in row.cells:\n"
            "        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER\n"
            "        for paragraph in cell.paragraphs:\n"
            "            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER\n"
            "            for run in paragraph.runs:\n"
            "                run.font.name = 'Times New Roman'\n"
            "                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')\n"
            "        if row_idx == 0:  # Header row\n"
            "            shading_elm = parse_xml(f'<w:shd {nsdecls} w:fill=\"1e40af\"/>')\n"
            "            cell._tc.get_or_add_tcPr().append(shading_elm)\n"
            "            for run in cell.paragraphs[0].runs:\n"
            "                run.font.color.rgb = RGBColor(255, 255, 255)\n"
            "        elif row_idx % 2 == 0:  # Even rows (zebra)\n"
            "            shading_elm = parse_xml(f'<w:shd {nsdecls} w:fill=\"f8f9fa\"/>')\n"
            "            cell._tc.get_or_add_tcPr().append(shading_elm)\n"
            "```\n\n"
            "**4. Body Text Format:**\n"
            "```python\n"
            "# Body text format\n"
            "paragraph_format = doc.styles['Normal'].paragraph_format\n"
            "paragraph_format.first_line_indent = Pt(24)  # 2 chars indent\n"
            "paragraph_format.line_spacing = 1.5\n"
            "```\n\n"
            f"{json.dumps(document_render_request, ensure_ascii=False)}"
        )

    @staticmethod
    def _extract_document_render_report(code_result: dict[str, Any]) -> dict[str, Any]:
        for item in list(code_result.get("tool_outputs") or []):
            if isinstance(item, dict):
                data = item.get("data") if isinstance(item.get("data"), dict) else {}
                render_report = data.get("render_report") if isinstance(data.get("render_report"), dict) else {}
                if render_report:
                    return render_report
        return {}

    def _generate_report(self, student: Student) -> dict[str, Any]:
        report = self.report_service.generate_report(student.id, None)
        pdf_path_text = str(report.get("pdf_path") or "").strip()
        if not pdf_path_text:
            return self._needs_input(
                task_type="generate_report",
                question="报告已生成，但未返回 PDF 路径，请稍后重试。",
                missing_fields=["report_pdf"],
            )
        pdf_path = Path(pdf_path_text)
        if not pdf_path.exists():
            return self._needs_input(
                task_type="generate_report",
                question="报告已生成，但未找到 PDF 文件，请稍后重试。",
                missing_fields=["report_pdf"],
            )

        attachment = self._register_generated_attachment(
            student_id=student.id,
            abs_path=pdf_path,
            file_name=f"{report.get('title') or 'career-report'}.pdf",
            description="agent generated report pdf",
        )
        summary = "已生成报告并提供可下载 PDF。"
        tool_output = self._tool_output(
            tool="generate_report",
            title="报告生成",
            summary=summary,
            data={
                "report_id": report.get("id"),
                "report_title": report.get("title"),
                "report_summary": report.get("summary"),
                "artifacts": [attachment],
            },
            next_actions=["查看报告详情", "继续优化行动计划"],
            card_type="report_card",
            context_patch={"context_binding": {"report_id": report.get("id"), "report_title": report.get("title")}},
        )
        return self._success(
            task_type="generate_report",
            reply=summary,
            tool_steps=[
                {"tool": "generate_report", "status": "done", "text": "done: 鐢熸垚鎶ュ憡"},
                {"tool": "export_report_pdf", "status": "done", "text": "done: 瀵煎嚭鎶ュ憡 PDF"},
            ],
            artifacts=[attachment],
            context_patch=tool_output.get("context_patch") or {},
            tool_outputs=[tool_output],
        )

    def _export_optimized_resume_from_context(
        self,
        student: Student,
        *,
        message: str | None = None,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        # 按优先级从多个上下文来源获取上一轮优化结果
        last_optimization = self._first_context_value(
            "last_resume_optimization",
            client_state,
            context_binding,
            (session_state.get("context_binding") if isinstance(session_state.get("context_binding"), dict) else {}),
            session_state,
        )

        # 如果找不到优化结果，提示用户先优化简历
        if not last_optimization or not isinstance(last_optimization, dict):
            return self._needs_input(
                task_type="export_optimized_resume",
                question="未找到上一次的简历优化结果，请先上传简历并执行'优化简历'操作。",
                missing_fields=["last_resume_optimization"],
            )

        # 提取优化结果中的关键数据
        optimization_result = last_optimization.get("optimization_result") if isinstance(last_optimization.get("optimization_result"), dict) else {}
        if not optimization_result:
            return self._needs_input(
                task_type="export_optimized_resume",
                question="优化结果数据不完整，请重新执行'优化简历'操作。",
                missing_fields=["optimization_result"],
            )

        optimized_document = optimization_result.get("optimized_resume_document")
        if not optimized_document:
            return self._needs_input(
                task_type="export_optimized_resume",
                question="优化后的简历文档内容为空，请重新执行'优化简历'操作。",
                missing_fields=["optimized_resume_document"],
            )

        try:
            # 准备输出路径
            now = datetime.now()
            output_dir = self.settings.upload_path / "resume_exports" / f"student_{student.id}"
            output_dir.mkdir(parents=True, exist_ok=True)

            source_file_name = str(last_optimization.get("source_file_name") or "resume")
            stem = Path(source_file_name).stem
            output_name = self._sanitize_display_filename(
                file_name=f"{stem}-optimized-resume.docx",
                default_name="optimized-resume.docx",
            )
            output_path = self._next_available_path(output_dir=output_dir, stem=f"{stem}-optimized", suffix=".docx")

            # 使用 ResumeRenderService 生成 Word 文档
            ResumeRenderService().render_word(
                resume_document=optimized_document,
                output_path=output_path,
            )

            # 注册生成的附件
            attachment = self._register_generated_attachment(
                student_id=student.id,
                abs_path=output_path,
                file_name=output_name,
                description="agent exported optimized resume docx from context",
            )

            # 构建成功响应
            summary = "已从上下文中导出上一版优化简历为 Word 文档。"
            tool_output = self._tool_output(
                tool="export_optimized_resume",
                title="导出优化简历",
                summary=summary,
                data={
                    "artifacts": [attachment],
                    "source_file_name": source_file_name,
                    "target_role": last_optimization.get("target_role") or optimization_result.get("target_role") or "",
                    "optimization_result": optimization_result,
                    "last_optimization_meta": {
                        "attachment_id": last_optimization.get("attachment_id"),
                        "created_at": last_optimization.get("created_at"),
                    },
                },
                next_actions=["继续生成职业规划报告", "继续岗位匹配分析"],
                card_type="resume_card",
                context_patch={
                    "context_binding": {
                        "last_exported_optimized_resume": {
                            "artifact_id": attachment.get("id"),
                            "file_name": output_name,
                            "exported_at": now.isoformat(timespec="seconds"),
                        }
                    }
                },
            )

            result = self._success(
                task_type="export_optimized_resume",
                reply=summary,
                tool_steps=[
                    {"tool": "resolve_last_optimization_context", "status": "done", "text": "done: resolved last resume optimization"},
                    {"tool": "render_optimized_resume_docx", "status": "done", "text": "done: rendered optimized resume to DOCX"},
                    {"tool": "register_artifact", "status": "done", "text": "done: registered DOCX artifact"},
                ],
                artifacts=[attachment],
                context_patch=tool_output.get("context_patch") or {},
                tool_outputs=[tool_output],
            )
            result["data"] = tool_output.get("data") or {}
            result["agent_flow_patch"] = [
                {"agent": "file", "action": "resolve_last_optimization_context"},
                {"agent": "file", "action": "render_optimized_resume_docx"},
                {"agent": "file", "action": "register_docx_artifact"},
                {"agent": "supervisor", "action": "check_exported_resume"},
            ]
            return result

        except Exception as exc:
            return self._failed(
                task_type="export_optimized_resume",
                failure_reason=f"导出优化简历失败：{str(exc)}",
                tool_steps=[
                    {"tool": "render_optimized_resume_docx", "status": "failed", "text": f"failed: {str(exc)}"},
                ],
            )

    def _generate_document(
        self,
        student: Student,
        message: str,
        *,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        from docx import Document

        # 检查是否有上一轮优化结果，如果有则委托给专用导出方法
        last_opt = self._first_context_value(
            "last_resume_optimization",
            client_state,
            context_binding,
            (session_state.get("context_binding") if isinstance(session_state.get("context_binding"), dict) else {}),
            session_state,
        )
        if last_opt and isinstance(last_opt, dict):
            return self._export_optimized_resume_from_context(
                student,
                message=message,
                session_state=session_state,
                context_binding=context_binding,
                client_state=client_state,
            )

        now = datetime.now()
        output_dir = self.settings.upload_path / "agent_files" / f"student_{student.id}"
        output_dir.mkdir(parents=True, exist_ok=True)

        code_bundle = self._resolve_generated_code_bundle(
            session_state=session_state,
            context_binding=context_binding,
            client_state=client_state,
        )
        bundle_files = self._extract_code_files(code_bundle)
        wants_code_package = self._looks_like_code_package_request(message)
        if wants_code_package and not bundle_files:
            return self._needs_input(
                task_type="generate_document",
                question="当前会话里没有可打包的代码结果。请先让 Code Agent 生成代码后再试。",
                missing_fields=["generated_code_bundle"],
            )

        file_meta = self._build_document_file_meta(
            message=message,
            generated_at=now,
            code_bundle=code_bundle if bundle_files else {},
        )
        output_path = self._next_available_path(output_dir=output_dir, stem=file_meta["path_stem"], suffix=".docx")

        document = Document()
        if bundle_files:
            self._render_code_package_document(
                document=document,
                message=message,
                generated_at=now,
                code_bundle=code_bundle,
                bundle_files=bundle_files,
            )
            summary = f"已生成代码打包 Word 文档《{file_meta['display_name']}》。"
            description = "agent generated code package document"
        else:
            # 没有有效内容可生成文档
            return self._needs_input(
                task_type="generate_document",
                question="请说明要生成什么内容，或发送'导出上一版优化简历为 Word'。",
                missing_fields=["document_content"],
            )

        document.save(str(output_path))

        attachment = self._register_generated_attachment(
            student_id=student.id,
            abs_path=output_path,
            file_name=file_meta["display_name"],
            description=description,
        )
        tool_data: dict[str, Any] = {
            "artifacts": [attachment],
            "document_meta": {
                "display_name": file_meta["display_name"],
                "path_stem": file_meta["path_stem"],
                "generated_at": now.strftime("%Y-%m-%d %H:%M:%S"),
                "source": "code_bundle" if bundle_files else "message_template",
            },
        }
        if bundle_files:
            tool_data["code_bundle"] = {
                "language": str(code_bundle.get("language") or ""),
                "run_id": str(code_bundle.get("run_id") or ""),
                "file_count": len(bundle_files),
                "verification_report": code_bundle.get("verification_report") or {},
            }
            tool_data["bundle_meta"] = {
                "included_file_paths": [row["path"] for row in bundle_files[:50]],
                "truncated_file_count": max(0, len(bundle_files) - 50),
            }

        tool_output = self._tool_output(
            tool="generate_document",
            title="文档生成",
            summary=summary,
            data=tool_data,
            next_actions=["继续生成图表", "继续生成图片"],
            card_type="action_checklist_card",
            context_patch={"context_binding": {"pending_file_offer": {}}},
        )
        return self._success(
            task_type="generate_document",
            reply=summary,
            tool_steps=[{"tool": "generate_document", "status": "done", "text": "done: 生成文档"}],
            artifacts=[attachment],
            context_patch={"context_binding": {"pending_file_offer": {}}},
            tool_outputs=[tool_output],
        )

    def _generate_chart(self, student: Student, message: str) -> dict[str, Any]:
        from PIL import Image, ImageDraw

        now = datetime.now()
        values = [int(item) for item in re.findall(r"\d+", str(message or ""))[:4]]
        if not values:
            values = [68, 82, 74, 90]
        labels = ["A", "B", "C", "D"][: len(values)]
        max_value = max(values) or 1

        output_dir = self.settings.upload_path / "agent_files" / f"student_{student.id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = now.strftime("%Y%m%d_%H%M%S")
        output_path = self._next_available_path(
            output_dir=output_dir,
            stem=f"chart-{now.strftime('%Y%m%d-%H%M%S')}",
            suffix=".png",
        )

        image = Image.new("RGB", (960, 540), (248, 250, 252))
        draw = ImageDraw.Draw(image)
        draw.rectangle((90, 60, 900, 460), outline=(148, 163, 184), width=2)

        bar_width = 120
        start_x = 160
        gap = 130
        for index, value in enumerate(values):
            bar_height = int((value / max_value) * 320)
            x1 = start_x + index * gap
            y1 = 440 - bar_height
            x2 = x1 + bar_width
            y2 = 440
            draw.rectangle((x1, y1, x2, y2), fill=(37, 99, 235))
            draw.text((x1 + 44, 448), labels[index], fill=(30, 41, 59))
            draw.text((x1 + 34, y1 - 20), str(value), fill=(15, 23, 42))

        draw.text((96, 24), "Agent Chart", fill=(15, 23, 42))
        image.save(str(output_path))

        attachment = self._register_generated_attachment(
            student_id=student.id,
            abs_path=output_path,
            file_name=f"图表_{timestamp}.png",
            description="agent generated chart image",
        )
        summary = "已生成可下载图表图片。"
        tool_output = self._tool_output(
            tool="generate_chart",
            title="图表生成",
            summary=summary,
            data={
                "artifacts": [attachment],
                "chart_meta": {
                    "type": "bar",
                    "labels": labels,
                    "values": values,
                    "max_value": max_value,
                    "canvas": {"width": 960, "height": 540},
                    "source": "message_numbers" if re.findall(r"\d+", str(message or "")) else "default_seed",
                },
            },
            next_actions=["继续生成图片", "继续生成可下载文档"],
            card_type="action_checklist_card",
            context_patch={},
        )
        return self._success(
            task_type="generate_chart",
            reply=summary,
            tool_steps=[{"tool": "generate_chart", "status": "done", "text": "done: 生成图表"}],
            artifacts=[attachment],
            context_patch={},
            tool_outputs=[tool_output],
        )

    def _generate_image(self, student: Student, message: str) -> dict[str, Any]:
        from PIL import Image, ImageDraw

        now = datetime.now()
        output_dir = self.settings.upload_path / "agent_files" / f"student_{student.id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = now.strftime("%Y%m%d_%H%M%S")
        output_path = self._next_available_path(
            output_dir=output_dir,
            stem=f"image-{now.strftime('%Y%m%d-%H%M%S')}",
            suffix=".png",
        )

        image = Image.new("RGB", (960, 540), (14, 165, 233))
        draw = ImageDraw.Draw(image)
        draw.ellipse((120, 100, 420, 400), fill=(37, 99, 235))
        draw.rectangle((520, 120, 860, 380), fill=(15, 23, 42))
        draw.text((88, 40), "Agent Image", fill=(255, 255, 255))
        draw.text((88, 470), (message or "generated by file agent")[:70], fill=(255, 255, 255))
        image.save(str(output_path))

        attachment = self._register_generated_attachment(
            student_id=student.id,
            abs_path=output_path,
            file_name=f"图片_{timestamp}.png",
            description="agent generated image",
        )
        summary = "已生成可下载图片。"
        tool_output = self._tool_output(
            tool="generate_image",
            title="图片生成",
            summary=summary,
            data={
                "artifacts": [attachment],
                "image_meta": {
                    "canvas": {"width": 960, "height": 540},
                    "prompt_excerpt": (message or "generated by file agent")[:120],
                    "render_shapes": ["ellipse", "rectangle", "text"],
                    "style": "flat_graphic",
                },
            },
            next_actions=["继续生成图表", "继续生成可下载文档"],
            card_type="action_checklist_card",
            context_patch={},
        )
        return self._success(
            task_type="generate_image",
            reply=summary,
            tool_steps=[{"tool": "generate_image", "status": "done", "text": "done: 生成图片"}],
            artifacts=[attachment],
            context_patch={},
            tool_outputs=[tool_output],
        )

    def _resolve_attachment(
        self,
        student_id: int,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> StudentAttachment | None:
        resolved = self._resolve_attachment_context(student_id, session_state, context_binding, client_state)
        attachment = resolved.get("attachment")
        return attachment if isinstance(attachment, StudentAttachment) else None

    def _resolve_attachment_context(
        self,
        student_id: int,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        state_binding = session_state.get("context_binding") if isinstance(session_state.get("context_binding"), dict) else {}
        candidate_ids = [
            client_state.get("attachment_id"),
            context_binding.get("attachment_id"),
            (context_binding.get("attachment") or {}).get("id") if isinstance(context_binding.get("attachment"), dict) else None,
            state_binding.get("attachment_id"),
            (state_binding.get("attachment") or {}).get("id") if isinstance(state_binding.get("attachment"), dict) else None,
            session_state.get("current_attachment_id"),
        ]
        saw_attachment_id = False
        for value in candidate_ids:
            attachment_id = self._to_int(value)
            if not attachment_id:
                continue
            saw_attachment_id = True
            row = self._attachment_by_id(student_id=student_id, attachment_id=attachment_id)
            if row:
                return {"attachment": row}

        resume_version_id = self._first_context_int(
            "resume_version_id",
            client_state=client_state,
            context_binding=context_binding,
            state_binding=state_binding,
            session_state=session_state,
        )
        if resume_version_id:
            row = self._attachment_from_resume_version(student_id=student_id, resume_version_id=resume_version_id)
            if row:
                return {"attachment": row}
            return self._invalid_attachment_resolution(
                invalid_fields=["resume_version_id"],
                question="当前选择的简历版本不存在或没有可用附件，请重新选择简历或上传新的简历附件。",
            )

        resume_id = self._first_context_int(
            "resume_id",
            client_state=client_state,
            context_binding=context_binding,
            state_binding=state_binding,
            session_state=session_state,
        )
        if resume_id:
            row = self._attachment_from_resume(student_id=student_id, resume_id=resume_id)
            if row:
                return {"attachment": row}
            return self._invalid_attachment_resolution(
                invalid_fields=["resume_id"],
                question="当前选择的简历不存在或没有可用附件，请重新选择简历或上传新的简历附件。",
            )

        if saw_attachment_id:
            return self._invalid_attachment_resolution(
                invalid_fields=["attachment_id"],
                question="当前选择的附件不存在或不可用，请重新上传简历附件后再试。",
            )

        return {"attachment": None}

    def _attachment_by_id(self, *, student_id: int, attachment_id: int | None) -> StudentAttachment | None:
        if not attachment_id:
            return None
        return (
            self.db.query(StudentAttachment)
            .filter(
                StudentAttachment.id == attachment_id,
                StudentAttachment.student_id == student_id,
                StudentAttachment.deleted.is_(False),
            )
            .first()
        )

    def _attachment_from_resume_version(self, *, student_id: int, resume_version_id: int) -> StudentAttachment | None:
        version = (
            self.db.query(StudentResumeVersion)
            .join(StudentResume, StudentResume.id == StudentResumeVersion.resume_id)
            .options(joinedload(StudentResumeVersion.attachment))
            .filter(
                StudentResumeVersion.id == resume_version_id,
                StudentResumeVersion.deleted.is_(False),
                StudentResume.student_id == student_id,
                StudentResume.deleted.is_(False),
            )
            .first()
        )
        if not version:
            return None
        return self._attachment_by_id(student_id=student_id, attachment_id=self._to_int(getattr(version, "attachment_id", None)))

    def _attachment_from_resume(self, *, student_id: int, resume_id: int) -> StudentAttachment | None:
        resume = (
            self.db.query(StudentResume)
            .options(joinedload(StudentResume.current_version), joinedload(StudentResume.source_attachment))
            .filter(
                StudentResume.id == resume_id,
                StudentResume.student_id == student_id,
                StudentResume.deleted.is_(False),
            )
            .first()
        )
        if not resume:
            return None
        current_version = getattr(resume, "current_version", None)
        current_attachment_id = None
        if current_version and not bool(getattr(current_version, "deleted", False)):
            current_attachment_id = self._to_int(getattr(current_version, "attachment_id", None))
        attachment = self._attachment_by_id(student_id=student_id, attachment_id=current_attachment_id)
        if attachment:
            return attachment
        return self._attachment_by_id(student_id=student_id, attachment_id=self._to_int(getattr(resume, "source_attachment_id", None)))

    def _first_context_int(
        self,
        field: str,
        *,
        client_state: dict[str, Any],
        context_binding: dict[str, Any],
        state_binding: dict[str, Any],
        session_state: dict[str, Any],
    ) -> int | None:
        for source in (client_state, context_binding, state_binding, session_state):
            if not isinstance(source, dict):
                continue
            parsed = self._to_int(source.get(field))
            if parsed:
                return parsed
            resume = source.get("resume") if isinstance(source.get("resume"), dict) else {}
            parsed = self._to_int(resume.get(field))
            if parsed:
                return parsed
        return None

    @staticmethod
    def _invalid_attachment_resolution(*, invalid_fields: list[str], question: str) -> dict[str, Any]:
        invalid_set = {str(item) for item in invalid_fields if str(item)}
        binding: dict[str, Any] = {}
        if "attachment_id" in invalid_set:
            binding["attachment_id"] = None
            binding["attachment"] = {}
        if invalid_set.intersection({"resume_id", "resume_version_id"}):
            binding["resume_id"] = None
            binding["resume_version_id"] = None
            binding["resume"] = {}
            binding["attachment_id"] = None
            binding["attachment"] = {}
        return {
            "attachment": None,
            "invalid_fields": sorted(invalid_set),
            "question": question,
            "context_patch": {"context_binding": binding} if binding else {},
        }

    def _resolve_resume_optimization_options(
        self,
        *,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        state_binding = session_state.get("context_binding") if isinstance(session_state.get("context_binding"), dict) else {}
        resume_from_client = client_state.get("resume") if isinstance(client_state.get("resume"), dict) else {}
        resume_from_binding = context_binding.get("resume") if isinstance(context_binding.get("resume"), dict) else {}
        resume_from_state = state_binding.get("resume") if isinstance(state_binding.get("resume"), dict) else {}

        target_role = self._first_text(
            client_state.get("target_role"),
            context_binding.get("target_role"),
            state_binding.get("target_role"),
            client_state.get("target_job"),
            context_binding.get("target_job"),
            state_binding.get("target_job"),
        )
        target_job_id = self._to_int(
            client_state.get("target_job_id")
            or context_binding.get("target_job_id")
            or state_binding.get("target_job_id")
            or resume_from_client.get("target_job_id")
            or resume_from_binding.get("target_job_id")
            or resume_from_state.get("target_job_id")
        )
        resume_id = self._to_int(
            client_state.get("resume_id")
            or context_binding.get("resume_id")
            or state_binding.get("resume_id")
            or resume_from_client.get("resume_id")
            or resume_from_binding.get("resume_id")
            or resume_from_state.get("resume_id")
        )
        resume_version_id = self._to_int(
            client_state.get("resume_version_id")
            or context_binding.get("resume_version_id")
            or state_binding.get("resume_version_id")
            or resume_from_client.get("resume_version_id")
            or resume_from_binding.get("resume_version_id")
            or resume_from_state.get("resume_version_id")
        )
        job_description = self._first_text(
            client_state.get("job_description"),
            context_binding.get("job_description"),
            state_binding.get("job_description"),
        )

        options: dict[str, Any] = {}
        if target_role:
            options["target_role"] = target_role
        if target_job_id:
            options["target_job_id"] = target_job_id
        if resume_id:
            options["resume_id"] = resume_id
        if resume_version_id:
            options["resume_version_id"] = resume_version_id
        if job_description:
            options["job_description"] = job_description
        return options

    def _build_file_parse_context_patch(
        self,
        *,
        attachment: StudentAttachment | None,
        status: str,
        reason: str,
    ) -> dict[str, Any]:
        attachment_id = int(getattr(attachment, "id", 0) or 0)
        attachment_name = str(getattr(attachment, "file_name", "") or "")
        attachment_type = str(getattr(attachment, "file_type", "") or "")

        binding: dict[str, Any] = {}
        if attachment_id:
            binding["attachment_id"] = attachment_id
            binding["attachment"] = {
                "id": attachment_id,
                "file_name": attachment_name,
                "file_type": attachment_type,
            }
        binding["file_parse_state"] = (
            {
                "status": status,
                "task": "parse_file",
                "reason": reason,
                "attachment_id": attachment_id or None,
            }
            if status
            else {}
        )
        return {"context_binding": binding}

    def _resolve_attachment_abs_path(self, attachment: StudentAttachment) -> Path | None:
        return resolve_upload_reference(
            upload_root=self.settings.upload_path,
            reference=attachment.file_path,
            must_exist=True,
        )

    def _register_generated_attachment(
        self,
        *,
        student_id: int,
        abs_path: Path,
        file_name: str,
        description: str,
    ) -> dict[str, Any]:
        upload_url = self._abs_path_to_upload_url(abs_path)
        resolved_file_name = self._sanitize_display_filename(file_name=file_name, default_name=abs_path.name)
        row = StudentAttachment(
            student_id=student_id,
            file_name=resolved_file_name,
            file_path=upload_url,
            file_type=abs_path.suffix.lstrip(".").lower(),
            description=description,
        )
        self.db.add(row)
        self.db.flush()
        return {
            "id": row.id,
            "name": row.file_name,
            "type": row.file_type or "",
            "download_url": row.file_path,
            "mime_type": self._mime_type_by_suffix(abs_path.suffix.lower()),
        }

    @staticmethod
    def _sanitize_display_filename(*, file_name: str, default_name: str) -> str:
        candidate = str(file_name or "").strip() or str(default_name or "").strip()
        candidate = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", candidate)
        candidate = re.sub(r"\s+", " ", candidate).strip(" .")
        return candidate or str(default_name or "download.bin")

    @staticmethod
    def _resolve_generated_code_bundle(
        *,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        sources = [
            context_binding.get("generated_code_bundle"),
            (session_state.get("context_binding") or {}).get("generated_code_bundle"),
            client_state.get("generated_code_bundle"),
        ]
        for item in sources:
            if isinstance(item, dict):
                return item
        return {}

    @staticmethod
    def _extract_code_files(bundle: dict[str, Any]) -> list[dict[str, str]]:
        rows = bundle.get("files") if isinstance(bundle.get("files"), list) else []
        files: list[dict[str, str]] = []
        for row in rows:
            if not isinstance(row, dict):
                continue
            path = str(row.get("path") or "").strip()
            content = str(row.get("content") or "")
            if not path or not content:
                continue
            files.append({"path": path, "content": content})
        return files

    @staticmethod
    def _looks_like_code_package_request(message: str) -> bool:
        compact = re.sub(r"\s+", "", str(message or "").lower())
        code_tokens = ("代码", "code", "源码", "脚本", "program", "script")
        package_tokens = ("打包", "下载", "文档", "word", "doc", "docx", "导出", "export", "package")
        return any(token in compact for token in code_tokens) and any(token in compact for token in package_tokens)

    def _build_document_file_meta(
        self,
        *,
        message: str,
        generated_at: datetime,
        code_bundle: dict[str, Any],
    ) -> dict[str, str]:
        display_stamp = generated_at.strftime("%Y%m%d_%H%M%S")
        path_stamp = generated_at.strftime("%Y%m%d-%H%M%S")
        if code_bundle:
            language = re.sub(r"[^a-z0-9]+", "", str(code_bundle.get("language") or "").lower()) or "code"
            return {
                "display_name": f"代码打包_{language}_{display_stamp}.docx",
                "path_stem": f"code-package-{language}-{path_stamp}",
            }

        title_hint = self._extract_message_title_hint(message)
        path_hint = self._slugify_ascii(title_hint) if title_hint else ""
        display_prefix = title_hint or "生成文档"
        path_prefix = path_hint or "generated-document"
        return {
            "display_name": f"{display_prefix}_{display_stamp}.docx",
            "path_stem": f"{path_prefix}-{path_stamp}",
        }

    @staticmethod
    def _extract_message_title_hint(message: str) -> str:
        normalized = str(message or "").strip()
        if not normalized:
            return ""
        cleaned = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]+", "", normalized)
        cleaned = cleaned.strip()
        return cleaned[:16] if cleaned else ""

    @staticmethod
    def _slugify_ascii(value: str) -> str:
        ascii_text = str(value or "").encode("ascii", "ignore").decode("ascii")
        slug = re.sub(r"[^a-zA-Z0-9]+", "-", ascii_text).strip("-").lower()
        return slug[:48]

    @staticmethod
    def _next_available_path(*, output_dir: Path, stem: str, suffix: str) -> Path:
        normalized_stem = re.sub(r"[^a-zA-Z0-9_-]+", "-", str(stem or "").strip()).strip("-") or "artifact"
        candidate = output_dir / f"{normalized_stem}{suffix}"
        index = 2
        while candidate.exists():
            candidate = output_dir / f"{normalized_stem}-{index}{suffix}"
            index += 1
        return candidate

    def _render_code_package_document(
        self,
        *,
        document: Any,
        message: str,
        generated_at: datetime,
        code_bundle: dict[str, Any],
        bundle_files: list[dict[str, str]],
    ) -> None:
        language = str(code_bundle.get("language") or "unknown")
        run_id = str(code_bundle.get("run_id") or "")
        verification = code_bundle.get("verification_report") or {}
        compile_summary = str((verification.get("compile") or {}).get("summary") or "compile/syntax checks passed")
        test_summary = str((verification.get("tests") or {}).get("summary") or "self-tests passed")

        document.add_heading("代码打包文档", level=1)
        document.add_paragraph(f"任务：{message or '未提供任务描述'}")
        document.add_paragraph(f"生成时间：{generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_paragraph(f"语言：{language}")
        if run_id:
            document.add_paragraph(f"Run ID：{run_id}")

        document.add_heading("校验摘要", level=2)
        document.add_paragraph(f"编译/语法：{compile_summary}")
        document.add_paragraph(f"自测结果：{test_summary}")

        document.add_heading("文件清单", level=2)
        for row in bundle_files:
            document.add_paragraph(row["path"], style="List Bullet")

        document.add_heading("源码内容", level=2)
        max_files = 20
        max_chars_per_file = 20000
        for index, row in enumerate(bundle_files):
            if index >= max_files:
                document.add_paragraph(f"其余 {len(bundle_files) - max_files} 个文件已省略。")
                break
            document.add_heading(row["path"], level=3)
            content = row["content"]
            truncated = False
            if len(content) > max_chars_per_file:
                content = f"{content[:max_chars_per_file]}\n\n... (内容过长，已截断)"
                truncated = True
            self._add_code_block_paragraph(document=document, content=content)
            if truncated:
                document.add_paragraph("提示：该文件内容过长，文档中仅保留前 20,000 个字符。")

    @staticmethod
    def _add_code_block_paragraph(*, document: Any, content: str) -> None:
        paragraph = document.add_paragraph(style="No Spacing")
        run = paragraph.add_run(content)
        run.font.name = "Consolas"

    def _export_optimized_resume_pdf(self, *, student_id: int, attachment_id: int, optimization: dict[str, Any]) -> Path:
        output_dir = self.settings.upload_path / "resume_exports" / f"student_{student_id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"student_{student_id}_attachment_{attachment_id}_optimized_resume.pdf"

        resume_document = optimization.get("optimized_resume_document") if isinstance(optimization, dict) else {}
        if not isinstance(resume_document, dict) or not resume_document:
            raise RuntimeError("optimized resume document is empty")

        ResumeRenderService().render_pdf(resume_document=resume_document, output_path=output_path)
        return output_path

    @staticmethod
    def _copy_as_stamped_file(*, source_path: Path, folder: Path, suffix: str) -> Path:
        folder.mkdir(parents=True, exist_ok=True)
        output = folder / f"agent_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}{suffix}"
        output.write_bytes(source_path.read_bytes())
        return output

    def _upload_url_to_abs_path(self, value: str) -> Path | None:
        return resolve_upload_reference(
            upload_root=self.settings.upload_path,
            reference=value,
            must_exist=True,
        )

    def _abs_path_to_upload_url(self, value: Path) -> str:
        return upload_path_to_url(upload_root=self.settings.upload_path, absolute_path=value)

    def _student_by_user(self, user_id: int) -> Student | None:
        return self.db.query(Student).filter(Student.user_id == user_id, Student.deleted.is_(False)).first()

    @staticmethod
    def _is_word_attachment(attachment: StudentAttachment) -> bool:
        file_type = str(attachment.file_type or "").strip().lower()
        if file_type in {"doc", "docx"}:
            return True
        suffix = Path(str(attachment.file_name or "")).suffix.lstrip(".").lower()
        return suffix in {"doc", "docx"}

    @staticmethod
    def _is_word_parse_failed(parsed: dict[str, Any]) -> bool:
        if bool(parsed.get("parser_failed")):
            return True
        if str(parsed.get("parser_failure_reason") or "") == "word_text_extraction_failed":
            return True
        try:
            return int(parsed.get("raw_text_length") or 0) <= 0
        except (TypeError, ValueError):
            return True

    @staticmethod
    def _mime_type_by_suffix(suffix: str) -> str:
        mapping = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
        }
        return mapping.get(suffix, "application/octet-stream")

    @staticmethod
    def _to_int(value: Any) -> int | None:
        try:
            return int(value)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _first_context_value(
        field: str,
        *sources: dict[str, Any],
    ) -> Any:
        # 按优先级顺序从多个上下文来源中查找指定字段值
        for source in sources:
            if not isinstance(source, dict):
                continue
            value = source.get(field)
            if value is not None and value != "" and value != {} and value != []:
                return value
        return None

    @staticmethod
    def _first_text(*values: Any) -> str:
        for value in values:
            text = str(value or "").strip()
            if text:
                return text
        return ""

    @staticmethod
    def _tool_output(
        *,
        tool: str,
        title: str,
        summary: str,
        data: dict[str, Any],
        next_actions: list[str],
        card_type: str,
        context_patch: dict[str, Any],
    ) -> dict[str, Any]:
        return {
            "tool": tool,
            "title": title,
            "summary": summary,
            "data": data,
            "card": {
                "type": card_type,
                "tool": tool,
                "title": title,
                "summary": summary,
                "data": data,
            },
            "next_actions": next_actions,
            "context_patch": context_patch,
        }

    @staticmethod
    def _success(
        *,
        task_type: str,
        reply: str,
        tool_steps: list[dict[str, Any]],
        artifacts: list[dict[str, Any]],
        context_patch: dict[str, Any],
        tool_outputs: list[dict[str, Any]],
    ) -> dict[str, Any]:
        return {
            "status": "success",
            "reply": reply,
            "tool_steps": tool_steps,
            "artifacts": artifacts,
            "context_patch": context_patch,
            "requires_user_input": False,
            "question": "",
            "file_task": {"type": task_type, "status": "success"},
            "tool_outputs": tool_outputs,
        }

    @staticmethod
    def _failed(
        *,
        task_type: str,
        failure_reason: str,
        tool_steps: list[dict[str, Any]],
        context_patch: dict[str, Any] | None = None,
        tool_outputs: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        reason = str(failure_reason or "文件任务失败。")
        return {
            "status": "failed",
            "reply": reason,
            "tool_steps": tool_steps,
            "artifacts": [],
            "context_patch": context_patch or {},
            "requires_user_input": False,
            "question": reason,
            "failure_reason": reason,
            "file_task": {"type": task_type, "status": "failed", "failure_reason": reason},
            "tool_outputs": tool_outputs or [],
        }

    @staticmethod
    def _needs_input(
        *,
        task_type: str,
        question: str,
        missing_fields: list[str],
        invalid_fields: list[str] | None = None,
        context_patch: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        file_task = {"type": task_type, "status": "needs_input", "missing_fields": missing_fields}
        if invalid_fields:
            file_task["invalid_fields"] = [str(item) for item in invalid_fields if str(item)]
        return {
            "status": "needs_input",
            "reply": "",
            "tool_steps": [],
            "artifacts": [],
            "context_patch": context_patch or {},
            "requires_user_input": True,
            "question": question,
            "file_task": file_task,
            "tool_outputs": [],
        }
