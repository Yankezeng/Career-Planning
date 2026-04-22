from __future__ import annotations

import hashlib
import io
import json
import re
import subprocess
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from uuid import uuid4

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.auth import User
from app.services.agent.common.agent_llm_profiles import build_agent_llm_service
from app.services.agent.code_agent.code_execution_runner import CodeExecutionRunner
from app.services.llm_service import LLMService


class CodeAgent:
    MAX_ATTEMPTS = 3
    CODE_TASK_KEYWORDS = (
        "write code",
        "generate code",
        "code agent",
        "coding",
        "program",
        "script",
        "python",
        "c++",
        "cpp",
        "javascript",
        " js ",
        "html",
        "css",
        "vue",
        "vbs",
        "vbscript",
        "mermaid",
        "mermain",
        "写代码",
        "写个程序",
        "编程",
        "代码",
        "脚本",
    )

    LANGUAGE_KEYWORDS: list[tuple[str, tuple[str, ...]]] = [
        ("python", ("python", " py ", "py脚本", "python脚本", "python code")),
        ("cpp", ("c++", " cpp ", "cxx", "c plus plus", "c++代码", "cpp代码")),
        ("c", (" c语言", "c language", " c code", "c代码")),
        ("javascript", ("javascript", " js ", "node", "js代码", "javascript代码")),
        ("html", ("html", "网页结构")),
        ("css", ("css", "样式表", "style sheet")),
        ("vue", ("vue", "vue3", "vue2", "sfc")),
        ("vbs", ("vbs", "vbscript")),
        ("mermaid", ("mermaid", "mermain", "流程图代码", "图表代码")),
    ]

    DEFAULT_FILE_NAME = {
        "python": "main.py",
        "c": "main.c",
        "cpp": "main.cpp",
        "javascript": "main.js",
        "html": "index.html",
        "css": "style.css",
        "vue": "App.vue",
        "vbs": "main.vbs",
        "mermaid": "diagram.mmd",
    }

    def __init__(
        self,
        db: Session | None = None,
        runner: CodeExecutionRunner | None = None,
        llm_service: LLMService | None = None,
    ):
        self.db = db
        self.settings = get_settings()
        self.llm_service = llm_service or build_agent_llm_service("code_agent")
        self.runner = runner or CodeExecutionRunner()

    def detect_task(self, *, message: str, selected_skill: str = "") -> dict[str, Any]:
        text = str(message or "").strip()
        lowered = f" {text.lower()} "
        skill = str(selected_skill or "").strip().lower()

        if skill == "code-agent":
            return {
                "is_code_task": True,
                "language": self._detect_language(lowered) or "python",
                "reason": "selected_skill",
            }

        if any(token in lowered for token in self.CODE_TASK_KEYWORDS):
            return {
                "is_code_task": True,
                "language": self._detect_language(lowered) or "python",
                "reason": "keyword_match",
            }

        return {"is_code_task": False, "language": "", "reason": "no_match"}

    def execute(
        self,
        *,
        user: User,
        message: str,
        selected_skill: str = "",
        session_state: dict[str, Any] | None = None,
        context_binding: dict[str, Any] | None = None,
        client_state: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        session_state = session_state or {}
        context_binding = context_binding or {}
        client_state = client_state or {}
        file_extract_request = self._resolve_file_extract_request(
            session_state=session_state,
            context_binding=context_binding,
            client_state=client_state,
        )
        if file_extract_request:
            return self._execute_file_extract(
                user=user,
                message=message,
                file_extract_request=file_extract_request,
            )

        document_render_request = self._resolve_document_render_request(
            session_state=session_state,
            context_binding=context_binding,
            client_state=client_state,
        )
        if document_render_request:
            return self._execute_document_render(
                user=user,
                message=message,
                document_render_request=document_render_request,
            )

        detection = self.detect_task(message=message, selected_skill=selected_skill)
        language = str(detection.get("language") or "python")
        run_id = f"code_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}"
        constraints = self._extract_constraints(message)

        attempt_reports: list[dict[str, Any]] = []
        previous_errors: list[str] = []

        for attempt in range(1, self.MAX_ATTEMPTS + 1):
            generated = self._generate_code_package(
                user=user,
                message=message,
                language=language,
                attempt=attempt,
                constraints=constraints,
                previous_errors=previous_errors,
            )
            files = list(generated.get("files") or [])
            tests = list(generated.get("tests") or [])

            verification_report = self.runner.verify(language=language, files=files, tests=tests)
            attempt_reports.append(
                {
                    "attempt": attempt,
                    "language": language,
                    "generation_source": generated.get("source") or "llm",
                    "verification_report": verification_report,
                }
            )

            if verification_report.get("ok"):
                return self._success_result(
                    user=user,
                    run_id=run_id,
                    language=language,
                    attempt=attempt,
                    files=files,
                    tests=tests,
                    verification_report=verification_report,
                    attempts=attempt_reports,
                    constraints=constraints,
                    original_message=message,
                )

            previous_errors.append(CodeExecutionRunner.build_verification_summary(verification_report))

        return self._failed_result(
            language=language,
            attempts=attempt_reports,
            message=message,
            constraints=constraints,
        )

    @staticmethod
    def _resolve_document_render_request(
        *,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        state_binding = session_state.get("context_binding") if isinstance(session_state.get("context_binding"), dict) else {}
        candidates = [
            context_binding.get("document_render_request"),
            client_state.get("document_render_request"),
            state_binding.get("document_render_request"),
        ]
        for item in candidates:
            if isinstance(item, dict):
                return item
        return {}

    @staticmethod
    def _resolve_file_extract_request(
        *,
        session_state: dict[str, Any],
        context_binding: dict[str, Any],
        client_state: dict[str, Any],
    ) -> dict[str, Any]:
        state_binding = session_state.get("context_binding") if isinstance(session_state.get("context_binding"), dict) else {}
        candidates = [
            context_binding.get("file_extract_request"),
            client_state.get("file_extract_request"),
            state_binding.get("file_extract_request"),
        ]
        for item in candidates:
            if isinstance(item, dict):
                return item
        return {}

    def _execute_file_extract(
        self,
        *,
        user: User,
        message: str,
        file_extract_request: dict[str, Any],
    ) -> dict[str, Any]:
        source_path = Path(str(file_extract_request.get("source_path") or ""))
        if not source_path.exists() or not source_path.is_file():
            return self._failed_file_extract_result(
                failure_reason="file_extract source_path does not exist",
                extraction_report={"request": file_extract_request},
                original_message=message,
            )

        student_id = int(file_extract_request.get("student_id") or getattr(user, "id", 0) or 0)
        run_id = f"file_extract_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}"
        run_dir = self.settings.upload_path / "agent_files" / f"student_{student_id}" / "file_extract" / run_id
        run_dir.mkdir(parents=True, exist_ok=True)

        request_path = run_dir / "request.json"
        report_path = run_dir / "report.json"
        normalized_request = {
            "source_path": str(source_path),
            "file_name": str(file_extract_request.get("file_name") or source_path.name),
            "file_type": str(file_extract_request.get("file_type") or source_path.suffix.lstrip(".")).lower(),
            "max_text_chars": int(file_extract_request.get("max_text_chars") or 60000),
        }
        request_path.write_text(json.dumps(normalized_request, ensure_ascii=False, indent=2), encoding="utf-8")

        try:
            extraction_report = self._extract_file_in_process(normalized_request)
            report_path.write_text(json.dumps(extraction_report, ensure_ascii=False, indent=2), encoding="utf-8")
            extract_report = {"returncode": 0, "stdout": "file extract completed", "stderr": "", "summary": "passed"}
        except Exception as exc:
            extract_report = {
                "returncode": 1,
                "stdout": "",
                "stderr": f"{exc.__class__.__name__}: {exc}",
                "summary": "failed",
            }
            return self._failed_file_extract_result(
                failure_reason=f"file extractor failed: {extract_report['stderr'] or extract_report['stdout']}",
                extraction_report={"run_id": run_id, "extract": extract_report, "run_dir": str(run_dir)},
                original_message=message,
            )
        if not report_path.exists():
            return self._failed_file_extract_result(
                failure_reason="file extractor did not create report.json",
                extraction_report={"run_id": run_id, "extract": extract_report, "run_dir": str(run_dir)},
                original_message=message,
            )

        try:
            persisted_report = json.loads(report_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, UnicodeDecodeError) as exc:
            return self._failed_file_extract_result(
                failure_reason=f"file extractor report is invalid JSON: {exc}",
                extraction_report={"run_id": run_id, "extract": extract_report, "run_dir": str(run_dir)},
                original_message=message,
            )
        extraction_report = persisted_report

        if not isinstance(extraction_report, dict):
            extraction_report = {"status": "failed", "failure_reason": "file extractor report is not an object"}
        extraction_report.update(
            {
                "run_id": run_id,
                "run_dir": str(run_dir),
                "request_path": str(request_path),
                "report_path": str(report_path),
                "extract": extract_report,
            }
        )
        if str(extraction_report.get("status") or "") != "success":
            return self._failed_file_extract_result(
                failure_reason=str(extraction_report.get("failure_reason") or "file extractor returned failed status"),
                extraction_report=extraction_report,
                original_message=message,
            )
        return self._success_file_extract_result(
            user=user,
            extraction_report=extraction_report,
            original_message=message,
        )

    def _execute_document_render(
        self,
        *,
        user: User,
        message: str,
        document_render_request: dict[str, Any],
    ) -> dict[str, Any]:
        spec = document_render_request["spec"]
        output_format = str(document_render_request["output_format"])
        output_name = Path(str(document_render_request["output_name"])).name
        student_id = int(document_render_request["student_id"])
        if output_format != "docx" or Path(output_name).suffix.lower() != ".docx":
            return self._failed_document_render_result(
                failure_reason="document_render only accepts output_format=docx and a .docx output_name",
                render_report={},
                original_message=message,
            )

        run_id = f"document_render_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}"
        run_dir = self.settings.upload_path / "agent_files" / f"student_{student_id}" / "document_render" / run_id
        tests_dir = run_dir / "tests"
        tests_dir.mkdir(parents=True, exist_ok=True)

        spec_path = run_dir / "spec.json"
        main_path = run_dir / "main.py"
        test_path = tests_dir / "test_render.py"
        output_path = run_dir / output_name

        spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
        main_path.write_text(self._build_document_renderer_code(), encoding="utf-8")
        test_path.write_text(self._build_document_renderer_test_code(), encoding="utf-8")

        compile_report = self._compile_python_file(main_path)
        if compile_report["returncode"] != 0:
            return self._failed_document_render_result(
                failure_reason=f"renderer compile failed: {compile_report['stderr'] or compile_report['stdout']}",
                render_report={"run_id": run_id, "compile": compile_report, "run_dir": str(run_dir)},
                original_message=message,
            )

        test_report = self._run_document_render_command([sys.executable, str(test_path)], cwd=run_dir)
        if test_report["returncode"] != 0:
            return self._failed_document_render_result(
                failure_reason=f"renderer test failed: {test_report['stderr'] or test_report['stdout']}",
                render_report={"run_id": run_id, "compile": compile_report, "tests": test_report, "run_dir": str(run_dir)},
                original_message=message,
            )

        render_command_report = self._run_document_render_command([sys.executable, str(main_path), str(spec_path), str(output_path)], cwd=run_dir)
        if render_command_report["returncode"] != 0:
            return self._failed_document_render_result(
                failure_reason=f"DOCX render command failed: {render_command_report['stderr'] or render_command_report['stdout']}",
                render_report={
                    "run_id": run_id,
                    "compile": compile_report,
                    "tests": test_report,
                    "render": render_command_report,
                    "run_dir": str(run_dir),
                },
                original_message=message,
            )

        final_check = self._verify_document_render_output(output_path=output_path)
        if not bool(final_check.get("ok")):
            return self._failed_document_render_result(
                failure_reason=str(final_check["failure_reason"]),
                render_report={
                    "run_id": run_id,
                    "compile": compile_report,
                    "tests": test_report,
                    "render": render_command_report,
                    "final_check": final_check,
                    "run_dir": str(run_dir),
                    "output_path": str(output_path),
                },
                original_message=message,
            )

        render_report = {
            "status": "success",
            "run_id": run_id,
            "run_dir": str(run_dir),
            "spec_path": str(spec_path),
            "main_py_path": str(main_path),
            "test_path": str(test_path),
            "output_path": str(output_path),
            "compile": compile_report,
            "tests": test_report,
            "render": render_command_report,
            "final_check": final_check,
        }
        artifact = {
            "name": output_name,
            "type": "document",
            "download_url": self._to_upload_url(output_path),
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "path": str(output_path),
        }
        return self._success_document_render_result(
            user=user,
            render_report=render_report,
            artifact=artifact,
            original_message=message,
        )

    @staticmethod
    def _compile_python_file(path: Path) -> dict[str, Any]:
        try:
            source = path.read_text(encoding="utf-8")
            compile(source, str(path), "exec")
        except Exception as exc:
            return {
                "returncode": 1,
                "stdout": "",
                "stderr": f"{exc.__class__.__name__}: {exc}",
                "summary": "failed",
            }
        return {
            "returncode": 0,
            "stdout": "",
            "stderr": "",
            "summary": "passed",
        }

    @staticmethod
    def _run_document_render_command(command: list[str], *, cwd: Path) -> dict[str, Any]:
        completed = subprocess.run(command, cwd=str(cwd), capture_output=True, text=True, check=False)
        return {
            "returncode": completed.returncode,
            "stdout": completed.stdout,
            "stderr": completed.stderr,
            "summary": "passed" if completed.returncode == 0 else "failed",
        }

    @staticmethod
    def _verify_document_render_output(*, output_path: Path) -> dict[str, Any]:
        from docx import Document

        if not output_path.exists():
            return {"ok": False, "failure_reason": "DOCX file does not exist"}
        size = output_path.stat().st_size
        if size <= 5 * 1024:
            return {"ok": False, "failure_reason": "DOCX file is not larger than 5KB", "size": size}
        document = Document(str(output_path))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.extend(paragraph.text for paragraph in cell.paragraphs)
        text = "\n".join(chunks)
        forbidden_titles = ("摘要卡片", "企业匹配分析", "职业规划路径", "修改建议表", "行动清单", "章节结构", "FileAgent")
        leaked_titles = [title for title in forbidden_titles if title in text]
        if leaked_titles:
            return {"ok": False, "failure_reason": f"DOCX contains report-only sections: {', '.join(leaked_titles)}", "size": size}
        resume_titles = ("教育背景", "专业技能", "实习经历", "项目经历", "证书资质", "荣誉证书", "竞赛经历")
        found_titles = [title for title in resume_titles if title in text]
        if len(found_titles) < 2:
            return {"ok": False, "failure_reason": "DOCX missing enough resume section titles", "size": size}
        return {"ok": True, "size": size, "resume_titles": found_titles}

    @classmethod
    def _extract_file_in_process(cls, request: dict[str, Any]) -> dict[str, Any]:
        path = Path(str(request.get("source_path") or ""))
        source_file_type = cls._detect_file_type(path, str(request.get("file_type") or ""))
        max_text_chars = int(request.get("max_text_chars") or 60000)
        if not path.exists() or not path.is_file():
            return {"status": "failed", "failure_reason": "source file does not exist", "source_file_type": source_file_type}

        suffix = path.suffix.lower()
        image_suffixes = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}
        if source_file_type == "docx" or suffix == ".docx":
            extracted = cls._extract_docx_in_process(path, max_text_chars)
            source_file_type = "docx"
        elif source_file_type == "pdf" or suffix == ".pdf":
            extracted = cls._extract_pdf_in_process(path, max_text_chars)
            source_file_type = "pdf"
        elif suffix in image_suffixes or source_file_type in {item.lstrip(".") for item in image_suffixes}:
            extracted = cls._extract_image_in_process(path)
            source_file_type = source_file_type or suffix.lstrip(".")
        else:
            extracted = cls._extract_plain_text_in_process(path, max_text_chars)

        metadata = dict(extracted.get("metadata") or {})
        metadata.update({"source_size_bytes": int(path.stat().st_size), "source_suffix": suffix})
        return {
            "status": "success",
            "source_file": {
                "name": str(request.get("file_name") or path.name),
                "path": str(path),
                "type": source_file_type,
                "suffix": suffix,
                "size": int(path.stat().st_size),
            },
            "source_file_type": source_file_type,
            "text": str(extracted.get("text") or ""),
            "tables": list(extracted.get("tables") or []),
            "images": list(extracted.get("images") or []),
            "charts": list(extracted.get("charts") or []),
            "metadata": metadata,
            "warnings": list(extracted.get("warnings") or []),
        }

    @staticmethod
    def _detect_file_type(path: Path, requested_type: str) -> str:
        value = str(requested_type or "").strip().lower().lstrip(".")
        return value or path.suffix.lower().lstrip(".")

    @staticmethod
    def _trim_extracted_text(text: str, max_chars: int) -> str:
        value = str(text or "")
        if max_chars > 0 and len(value) > max_chars:
            return value[:max_chars] + "\n\n... (text truncated by extractor)"
        return value

    @classmethod
    def _extract_docx_in_process(cls, path: Path, max_text_chars: int) -> dict[str, Any]:
        warnings: list[str] = []
        paragraphs: list[str] = []
        tables: list[dict[str, Any]] = []
        try:
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
            for table_index, table in enumerate(document.tables, start=1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(cells)
                if rows:
                    tables.append({"index": table_index, "rows": rows})
        except Exception as exc:
            warnings.append(f"python-docx extraction failed, used DOCX XML fallback: {exc}")
            paragraphs = cls._extract_docx_xml_text(path)
            tables = cls._extract_docx_xml_tables(path)

        images, charts, scan_warnings = cls._scan_docx_media_and_charts(path)
        warnings.extend(scan_warnings)
        text = cls._trim_extracted_text("\n".join(paragraphs), max_text_chars)
        return {
            "text": text,
            "tables": tables,
            "images": images,
            "charts": charts,
            "metadata": {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "image_count": len(images),
                "chart_count": len(charts),
                "text_chars": len(text),
            },
            "warnings": warnings,
        }

    @staticmethod
    def _scan_docx_media_and_charts(path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[str]]:
        warnings: list[str] = []
        images: list[dict[str, Any]] = []
        charts: list[dict[str, Any]] = []
        try:
            with zipfile.ZipFile(path) as archive:
                for info in archive.infolist():
                    lower_name = info.filename.lower()
                    if lower_name.startswith("word/media/") and not info.is_dir():
                        image_info: dict[str, Any] = {"name": Path(info.filename).name, "path": info.filename, "bytes": int(info.file_size)}
                        image_info.update(CodeAgent._image_dimensions_from_bytes(archive.read(info.filename)))
                        images.append(image_info)
                    elif lower_name.startswith("word/charts/") and not info.is_dir():
                        charts.append({"name": Path(info.filename).name, "path": info.filename, "bytes": int(info.file_size)})
        except Exception as exc:
            warnings.append(f"DOCX zip media/chart scan failed: {exc}")
        return images, charts, warnings

    @staticmethod
    def _image_dimensions_from_bytes(data: bytes) -> dict[str, Any]:
        try:
            from PIL import Image

            with Image.open(io.BytesIO(data)) as image:
                return {"width": int(image.width), "height": int(image.height), "mode": str(image.mode)}
        except Exception:
            return {}

    @staticmethod
    def _extract_docx_xml_text(path: Path) -> list[str]:
        try:
            with zipfile.ZipFile(path) as archive:
                xml_bytes = archive.read("word/document.xml")
            root = ET.fromstring(xml_bytes)
        except Exception:
            return []
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        texts: list[str] = []
        current: list[str] = []
        for element in root.iter():
            if element.tag == f"{namespace}t" and element.text:
                current.append(element.text)
            elif element.tag == f"{namespace}p":
                paragraph = "".join(current).strip()
                if paragraph:
                    texts.append(paragraph)
                current = []
        paragraph = "".join(current).strip()
        if paragraph:
            texts.append(paragraph)
        return texts

    @staticmethod
    def _extract_docx_xml_tables(path: Path) -> list[dict[str, Any]]:
        try:
            with zipfile.ZipFile(path) as archive:
                xml_bytes = archive.read("word/document.xml")
            root = ET.fromstring(xml_bytes)
        except Exception:
            return []
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        tables: list[dict[str, Any]] = []
        for table_index, table in enumerate(root.iter(f"{namespace}tbl"), start=1):
            rows = []
            for row in table.iter(f"{namespace}tr"):
                cells = []
                for cell in row.iter(f"{namespace}tc"):
                    texts = [node.text or "" for node in cell.iter(f"{namespace}t") if node.text]
                    cells.append("".join(texts).strip())
                if any(cells):
                    rows.append(cells)
            if rows:
                tables.append({"index": table_index, "rows": rows})
        return tables

    @classmethod
    def _extract_pdf_in_process(cls, path: Path, max_text_chars: int) -> dict[str, Any]:
        warnings: list[str] = []
        text_parts: list[str] = []
        images: list[dict[str, Any]] = []
        try:
            import fitz

            with fitz.open(str(path)) as document:
                for page_index, page in enumerate(document, start=1):
                    page_text = page.get_text("text") or ""
                    if page_text.strip():
                        text_parts.append(page_text.strip())
                    for image_index, image in enumerate(page.get_images(full=True), start=1):
                        images.append(
                            {
                                "page": page_index,
                                "index": image_index,
                                "xref": int(image[0]),
                                "width": int(image[2]),
                                "height": int(image[3]),
                                "bits_per_component": image[4] if len(image) > 4 else None,
                            }
                        )
        except Exception as exc:
            warnings.append(f"PDF extraction requires PyMuPDF/fitz and failed: {exc}")
        text = cls._trim_extracted_text("\n".join(text_parts), max_text_chars)
        return {
            "text": text,
            "tables": [],
            "images": images,
            "charts": [],
            "metadata": {"page_count": len(text_parts), "image_count": len(images), "chart_count": 0, "text_chars": len(text)},
            "warnings": warnings,
        }

    @staticmethod
    def _extract_image_in_process(path: Path) -> dict[str, Any]:
        warnings: list[str] = []
        images: list[dict[str, Any]] = [{"name": path.name, "path": str(path), "bytes": int(path.stat().st_size)}]
        try:
            from PIL import Image

            with Image.open(str(path)) as image:
                images[0].update({"width": int(image.width), "height": int(image.height), "mode": str(image.mode)})
        except Exception as exc:
            warnings.append(f"image metadata extraction failed: {exc}")

        text = ""
        try:
            from PIL import Image
            import pytesseract

            with Image.open(str(path)) as image:
                text = str(pytesseract.image_to_string(image) or "").strip()
        except Exception as exc:
            warnings.append(f"local OCR unavailable or failed: {exc}")
        return {
            "text": text,
            "tables": [],
            "images": images,
            "charts": [],
            "metadata": {"image_count": 1, "chart_count": 0, "text_chars": len(text)},
            "warnings": warnings,
        }

    @classmethod
    def _extract_plain_text_in_process(cls, path: Path, max_text_chars: int) -> dict[str, Any]:
        warnings: list[str] = []
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception as exc:
            warnings.append(f"plain text read failed: {exc}")
            text = ""
        text = cls._trim_extracted_text(text, max_text_chars)
        return {
            "text": text,
            "tables": [],
            "images": [],
            "charts": [],
            "metadata": {"text_chars": len(text), "image_count": 0, "chart_count": 0},
            "warnings": warnings,
        }

    @staticmethod
    def _build_document_renderer_code() -> str:
        return r'''
from __future__ import annotations

import json
import math
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PRIMARY_BLUE = RGBColor(143, 175, 218)
LIGHT_BLUE = RGBColor(232, 240, 250)
PALE_BLUE = "DCE8F6"
BLUE_FILL = "8FAFDA"
LIGHT_FILL = "EEF4FB"
TEXT_DARK = RGBColor(31, 45, 61)
TEXT_MUTED = RGBColor(89, 105, 122)
WHITE = RGBColor(255, 255, 255)

SECTION_TITLES = ("教育背景", "专业技能", "核心技能", "实习经历", "项目经历", "证书资质", "荣誉证书", "竞赛经历", "校园经历")
REPORT_ONLY_TITLES = ("摘要卡片", "企业匹配分析", "职业规划路径", "修改建议表", "行动清单", "章节结构", "FileAgent")
DATE_RE = re.compile(r"\d{4}\.\d{1,2}\s*[-—~至]\s*(?:\d{4}\.\d{1,2}|至今|Present|Now)", re.I)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = "FFFFFF", size: str = "0", val: str = "nil") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_no_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.4, color: RGBColor = TEXT_DARK, alignment=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text or ""))
    set_run_font(run, bold=bold, size=size, color=color)


def set_run_font(run, *, bold: bool = False, size: float = 9.5, color: RGBColor = TEXT_DARK) -> None:
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = color


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(9.3)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(1.2)


def clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def split_lines(value: Any) -> list[str]:
    return [line.rstrip() for line in str(value or "").splitlines() if line.strip()]


def normalize_str_list(value: Any) -> list[str]:
    if not value:
        return []
    if isinstance(value, str):
        return [item.strip() for item in re.split(r"[\n,，、;；]+", value) if item.strip()]
    if isinstance(value, list):
        rows = []
        for item in value:
            if isinstance(item, dict):
                text = " | ".join(str(v).strip() for v in item.values() if str(v or "").strip())
            else:
                text = str(item or "").strip()
            if text:
                rows.append(text)
        return rows
    return [str(value).strip()]


def split_bullet_text(value: Any) -> list[str]:
    if isinstance(value, list):
        rows = []
        for item in value:
            rows.extend(split_bullet_text(item))
        return rows
    text = str(value or "").strip()
    if not text:
        return []
    rows: list[str] = []
    for line in text.splitlines():
        line = line.strip().lstrip("-• ").strip()
        if not line:
            continue
        parts = [item.strip(" ；;。") for item in re.split(r"[；;]\s*", line) if item.strip(" ；;。")]
        rows.extend(parts or [line])
    return rows


def join_non_empty(*values: Any, sep: str = " | ") -> str:
    return sep.join(str(value).strip() for value in values if str(value or "").strip())


def parse_contact_line(line: str, resume: dict[str, Any]) -> None:
    pairs = {
        "phone": r"(?:电话|手机)[:：]\s*([^|｜]+)",
        "email": r"(?:邮箱|Email|E-mail)[:：]\s*([^|｜]+)",
        "github": r"(?:GitHub|Github)[:：]\s*([^|｜]+)",
        "target_city": r"(?:所在地|现居地|城市)[:：]\s*([^|｜]+)",
    }
    for key, pattern in pairs.items():
        match = re.search(pattern, line, re.I)
        if match:
            resume[key] = match.group(1).strip()


def parse_plain_resume(text: str, spec: dict[str, Any]) -> dict[str, Any]:
    raw_lines = split_lines(str(text or "").replace("\r\n", "\n"))
    if raw_lines and raw_lines[0].startswith("正文"):
        raw_lines[0] = raw_lines[0].split(":", 1)[-1].split("：", 1)[-1].strip()
    resume: dict[str, Any] = {
        "name": "",
        "target_role": clean_text(spec.get("target_role")),
        "raw_sections": {},
    }
    current_title = ""
    for raw_line in raw_lines:
        line = raw_line.strip()
        title = line.strip("：:")
        if title in SECTION_TITLES:
            current_title = "专业技能" if title == "核心技能" else title
            resume["raw_sections"].setdefault(current_title, [])
            continue
        if not current_title:
            if line.startswith("求职意向"):
                resume["target_role"] = line.split("：", 1)[-1].split(":", 1)[-1].strip()
            elif any(marker in line for marker in ("电话", "邮箱", "GitHub", "Github")):
                parse_contact_line(line, resume)
            elif not resume["name"] and len(line) <= 12:
                resume["name"] = line
            continue
        resume["raw_sections"].setdefault(current_title, []).append(raw_line)
    resume["name"] = resume["name"] or clean_text(spec.get("name")) or "个人简历"
    return resume


def normalize_resume(spec: dict[str, Any]) -> dict[str, Any]:
    value = spec.get("optimized_resume_document") or spec.get("optimized_resume") or {}
    if isinstance(value, str):
        return parse_plain_resume(value, spec)
    if isinstance(value, dict):
        text_value = value.get("正文") or value.get("resume") or value.get("content")
        if isinstance(text_value, str) and not any(key in value for key in ("name", "skills", "projects", "internships")):
            return parse_plain_resume(text_value, spec)
        resume = dict(value)
        resume.setdefault("target_role", clean_text(spec.get("target_role")))
        resume.setdefault("raw_sections", {})
        resume["name"] = clean_text(resume.get("name")) or "个人简历"
        return resume
    return {"name": "个人简历", "target_role": clean_text(spec.get("target_role")), "raw_sections": {}}


def add_top_ribbon(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_no_borders(table)
    widths = [Inches(0.28), Inches(1.25), Inches(5.75)]
    fills = ["EEF2F7", "DEE7F2", BLUE_FILL]
    row = table.rows[0]
    row.height = Inches(0.16)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for index, cell in enumerate(row.cells):
        cell.width = widths[index]
        set_cell_shading(cell, fills[index])
        set_cell_text(cell, "")


def add_spacer(document: Document, size: float = 3) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(size)
    paragraph.paragraph_format.line_spacing = 0.5


def build_meta_items(resume: dict[str, Any]) -> list[tuple[str, str]]:
    school = join_non_empty(resume.get("college"), resume.get("major"))
    items = [
        ("电话", resume.get("phone")),
        ("邮箱", resume.get("email")),
        ("求职意向", resume.get("target_role")),
        ("GitHub", resume.get("github")),
        ("学校", school),
        ("所在地", resume.get("target_city")),
        ("学历", resume.get("grade")),
    ]
    return [(label, clean_text(value)) for label, value in items if clean_text(value)]


def add_header(document: Document, resume: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_no_borders(table)
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(5.75)
    photo_cell, info_cell = table.rows[0].cells
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    info_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    photo_table = photo_cell.add_table(rows=1, cols=1)
    set_table_no_borders(photo_table)
    photo_row = photo_table.rows[0]
    photo_row.height = Inches(1.35)
    photo_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    frame = photo_row.cells[0]
    frame.width = Inches(1.18)
    set_cell_shading(frame, "FFFFFF")
    set_cell_border(frame, color="D7DEE8", size="8", val="single")
    set_cell_text(frame, "")

    name_paragraph = info_cell.paragraphs[0]
    name_paragraph.paragraph_format.space_after = Pt(4)
    name_run = name_paragraph.add_run(clean_text(resume.get("name")) or "个人简历")
    set_run_font(name_run, bold=True, size=26, color=PRIMARY_BLUE)

    role = clean_text(resume.get("target_role"))
    if role:
        role_paragraph = info_cell.add_paragraph()
        role_paragraph.paragraph_format.space_after = Pt(7)
        role_run = role_paragraph.add_run(role)
        set_run_font(role_run, bold=True, size=10.5, color=TEXT_MUTED)

    items = build_meta_items(resume)
    if items:
        rows = math.ceil(len(items) / 2)
        meta_table = info_cell.add_table(rows=rows, cols=2)
        set_table_no_borders(meta_table)
        for index, (label, value) in enumerate(items):
            row = meta_table.rows[index // 2]
            cell = row.cells[index % 2]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            label_run = paragraph.add_run(f"{label}：")
            set_run_font(label_run, bold=True, size=9.5, color=TEXT_DARK)
            value_run = paragraph.add_run(value)
            set_run_font(value_run, size=9.5, color=TEXT_DARK)

    add_spacer(document, 4)


def add_section_heading(document: Document, title: str) -> None:
    add_spacer(document, 2)
    table = document.add_table(rows=1, cols=2)
    set_table_no_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(1.42)
    table.columns[1].width = Inches(5.8)
    row = table.rows[0]
    row.height = Inches(0.27)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    label_cell, tail_cell = row.cells
    set_cell_shading(label_cell, BLUE_FILL)
    set_cell_text(label_cell, f"  {title}", bold=True, size=11.2, color=WHITE)
    set_cell_text(tail_cell, "")


def add_body(document: Document, text: str, *, bold: bool = False, size: float = 9.3) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text or "").strip())
    set_run_font(run, bold=bold, size=size, color=TEXT_DARK)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1.1)
    paragraph.paragraph_format.line_spacing = 1.04
    run = paragraph.add_run(str(text or "").strip())
    set_run_font(run, size=9.15, color=TEXT_DARK)


def split_title_date(text: str) -> tuple[str, str]:
    match = DATE_RE.search(text)
    if not match:
        return text, ""
    date = match.group(0).strip()
    left = (text[: match.start()] + text[match.end() :]).strip(" |｜-—")
    return left or text, date


def add_meta_row(document: Document, left: str, right: str = "") -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_no_borders(table)
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(2.05)
    left_cell, right_cell = table.rows[0].cells
    set_cell_text(left_cell, left, bold=True, size=10.2)
    set_cell_text(right_cell, right, bold=True, size=10.2, alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_section_lines(document: Document, title: str, lines: list[str]) -> None:
    if not lines:
        return
    add_section_heading(document, title)
    for raw_line in lines:
        line = str(raw_line or "").strip()
        if not line:
            continue
        bullet_match = re.match(r"^\s*[-•]\s*(.+)$", line)
        if bullet_match:
            body = bullet_match.group(1).strip()
            if body.rstrip("：:") in {"项目成果", "核心职责", "项目描述", "技术栈"}:
                add_body(document, body, bold=True)
            else:
                add_bullet(document, body)
            continue
        if DATE_RE.search(line) or " | " in line or "｜" in line:
            left, right = split_title_date(line)
            add_meta_row(document, left, right)
        else:
            add_body(document, line)


def add_structured_experience(document: Document, title: str, items: Any, *, name_key: str, role_key: str, detail_key: str) -> None:
    if not items:
        return
    if isinstance(items, str):
        add_section_lines(document, title, split_lines(items))
        return
    add_section_heading(document, title)
    for item in items:
        if not isinstance(item, dict):
            add_body(document, str(item), bold=True)
            continue
        heading = join_non_empty(item.get(name_key), item.get(role_key))
        date = clean_text(item.get("duration"))
        if heading or date:
            add_meta_row(document, heading, date)
        technologies = normalize_str_list(item.get("technologies") or item.get("skills"))
        if technologies:
            add_bullet(document, f"技术栈：{'、'.join(technologies)}")
        for key in (detail_key, "description", "responsibilities", "results", "highlights"):
            for bullet in split_bullet_text(item.get(key)):
                add_bullet(document, bullet)


def add_structured_resume(document: Document, resume: dict[str, Any]) -> None:
    education = split_lines(resume.get("education_experience"))
    if not education and any(clean_text(resume.get(key)) for key in ("college", "major", "grade")):
        education = [join_non_empty(resume.get("college"), resume.get("major"), resume.get("grade"))]
    add_section_lines(document, "教育背景", education)

    skills = normalize_str_list(resume.get("skills"))
    if skills:
        add_section_heading(document, "专业技能")
        for skill in skills:
            add_bullet(document, skill)

    add_structured_experience(document, "实习经历", resume.get("internships"), name_key="company", role_key="position", detail_key="rewrite")
    add_structured_experience(document, "项目经历", resume.get("projects"), name_key="name", role_key="role", detail_key="rewrite")

    certificate_rows = normalize_str_list(resume.get("certificates"))
    competitions = resume.get("competitions") or []
    for item in competitions:
        if isinstance(item, dict):
            certificate_rows.append(join_non_empty(item.get("name"), item.get("award"), item.get("level"), item.get("description")))
        else:
            certificate_rows.append(str(item))
    if certificate_rows:
        add_section_heading(document, "证书资质")
        for item in certificate_rows:
            add_bullet(document, item)

    add_structured_experience(document, "校园经历", resume.get("campus_experiences"), name_key="title", role_key="role", detail_key="description")


def add_plain_resume(document: Document, resume: dict[str, Any]) -> None:
    sections = resume.get("raw_sections") or {}
    ordered = [
        ("教育背景", ("教育背景",)),
        ("专业技能", ("专业技能", "核心技能")),
        ("实习经历", ("实习经历",)),
        ("项目经历", ("项目经历",)),
        ("证书资质", ("证书资质", "荣誉证书", "竞赛经历")),
        ("校园经历", ("校园经历",)),
    ]
    for output_title, aliases in ordered:
        lines: list[str] = []
        for alias in aliases:
            lines = list(sections.get(alias) or [])
            if lines:
                break
        add_section_lines(document, output_title, lines)


def collect_document_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(chunks)


def render_document(spec: dict, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    resume = normalize_resume(spec)
    add_top_ribbon(document)
    add_spacer(document, 8)
    add_header(document, resume)
    if resume.get("raw_sections"):
        add_plain_resume(document, resume)
    else:
        add_structured_resume(document, resume)

    text = collect_document_text(document)
    leaked_titles = [title for title in REPORT_ONLY_TITLES if title in text]
    if leaked_titles:
        raise ValueError(f"report-only sections leaked into resume: {', '.join(leaked_titles)}")
    found_resume_titles = [title for title in ("教育背景", "专业技能", "实习经历", "项目经历", "证书资质") if title in text]
    if len(found_resume_titles) < 2:
        raise ValueError("rendered resume does not contain enough resume sections")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    spec_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    spec = json.loads(spec_path.read_text(encoding="utf-8"))
    render_document(spec, output_path)


if __name__ == "__main__":
    main()
'''.lstrip()

    @staticmethod
    def _build_document_renderer_test_code() -> str:
        return r'''
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from main import render_document


spec = json.loads((ROOT / "spec.json").read_text(encoding="utf-8"))
output_path = ROOT / "test-render.docx"
render_document(spec, output_path)

assert output_path.exists(), "DOCX file was not created"
assert output_path.stat().st_size > 5 * 1024, "DOCX file is smaller than 5KB"

cli_output_path = ROOT / "test-render-cli.docx"
completed = subprocess.run(
    [sys.executable, str(ROOT / "main.py"), str(ROOT / "spec.json"), str(cli_output_path)],
    cwd=str(ROOT),
    capture_output=True,
    text=True,
    check=False,
)
assert completed.returncode == 0, f"CLI render failed: {completed.stderr or completed.stdout}"
assert cli_output_path.exists(), "CLI DOCX file was not created"
assert cli_output_path.stat().st_size > 5 * 1024, "CLI DOCX file is smaller than 5KB"

document = Document(str(output_path))
text = "\n".join(
    [paragraph.text for paragraph in document.paragraphs]
    + [paragraph.text for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]
)
for forbidden in ("摘要卡片", "企业匹配分析", "职业规划路径", "修改建议表", "行动清单", "章节结构", "FileAgent"):
    assert forbidden not in text, f"forbidden report-only content leaked: {forbidden}"
found_titles = [title for title in ("教育背景", "专业技能", "实习经历", "项目经历", "证书资质") if title in text]
assert len(found_titles) >= 2, f"not enough resume section titles: {found_titles}"

print("document render self-test passed")
'''.lstrip()

    @staticmethod
    def _build_file_extractor_code() -> str:
        return r'''
from __future__ import annotations

import io
import json
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}


def detect_type(path: Path, requested_type: str) -> str:
    value = str(requested_type or "").strip().lower().lstrip(".")
    if value:
        return value
    return path.suffix.lower().lstrip(".")


def trim_text(text: str, max_chars: int) -> str:
    value = str(text or "")
    if max_chars > 0 and len(value) > max_chars:
        return value[:max_chars] + "\n\n... (text truncated by extractor)"
    return value


def image_dimensions(data: bytes) -> dict:
    try:
        from PIL import Image

        with Image.open(io.BytesIO(data)) as image:
            return {"width": int(image.width), "height": int(image.height), "mode": str(image.mode)}
    except Exception:
        return {}


def extract_docx(path: Path, max_text_chars: int) -> dict:
    warnings = []
    paragraphs = []
    tables = []

    try:
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        for table_index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                tables.append({"index": table_index, "rows": rows})
    except Exception as exc:
        raise RuntimeError(f"python-docx extraction failed: {exc}") from exc

    images = []
    charts = []
    try:
        with zipfile.ZipFile(path) as archive:
            for info in archive.infolist():
                lower_name = info.filename.lower()
                if lower_name.startswith("word/media/") and not info.is_dir():
                    image_info = {
                        "name": Path(info.filename).name,
                        "path": info.filename,
                        "bytes": int(info.file_size),
                    }
                    try:
                        image_info.update(image_dimensions(archive.read(info.filename)))
                    except Exception:
                        pass
                    images.append(image_info)
                elif lower_name.startswith("word/charts/") and not info.is_dir():
                    charts.append(
                        {
                            "name": Path(info.filename).name,
                            "path": info.filename,
                            "bytes": int(info.file_size),
                        }
                    )
    except Exception as exc:
        warnings.append(f"DOCX zip media/chart scan failed: {exc}")

    text = trim_text("\n".join(paragraphs), max_text_chars)
    metadata = {
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": len(images),
        "chart_count": len(charts),
        "text_chars": len(text),
    }
    return {
        "text": text,
        "tables": tables,
        "images": images,
        "charts": charts,
        "metadata": metadata,
        "warnings": warnings,
    }


def extract_docx_xml_text(path: Path) -> list[str]:
    try:
        with zipfile.ZipFile(path) as archive:
            xml_bytes = archive.read("word/document.xml")
    except Exception:
        return []

    try:
        root = ET.fromstring(xml_bytes)
    except Exception:
        return []
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    texts = []
    current = []
    for element in root.iter():
        if element.tag == f"{namespace}t" and element.text:
            current.append(element.text)
        elif element.tag == f"{namespace}p":
            paragraph = "".join(current).strip()
            if paragraph:
                texts.append(paragraph)
            current = []
    if current:
        paragraph = "".join(current).strip()
        if paragraph:
            texts.append(paragraph)
    return texts


def extract_docx_xml_tables(path: Path) -> list[dict]:
    try:
        with zipfile.ZipFile(path) as archive:
            xml_bytes = archive.read("word/document.xml")
    except Exception:
        return []

    try:
        root = ET.fromstring(xml_bytes)
    except Exception:
        return []

    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    tables = []
    for table_index, table in enumerate(root.iter(f"{namespace}tbl"), start=1):
        rows = []
        for row in table.iter(f"{namespace}tr"):
            cells = []
            for cell in row.iter(f"{namespace}tc"):
                texts = [node.text or "" for node in cell.iter(f"{namespace}t") if node.text]
                cell_text = "".join(texts).strip()
                cells.append(cell_text)
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"index": table_index, "rows": rows})
    return tables


def extract_pdf(path: Path, max_text_chars: int) -> dict:
    warnings = []
    text_parts = []
    images = []
    try:
        import fitz

        with fitz.open(str(path)) as document:
            for page_index, page in enumerate(document, start=1):
                page_text = page.get_text("text") or ""
                if page_text.strip():
                    text_parts.append(page_text.strip())
                for image_index, image in enumerate(page.get_images(full=True), start=1):
                    images.append(
                        {
                            "page": page_index,
                            "index": image_index,
                            "xref": int(image[0]),
                            "width": int(image[2]),
                            "height": int(image[3]),
                            "bits_per_component": int(image[4]) if len(image) > 4 and str(image[4]).isdigit() else image[4],
                        }
                    )
    except Exception as exc:
        warnings.append(f"PDF extraction requires PyMuPDF/fitz and failed: {exc}")
    text = trim_text("\n".join(text_parts), max_text_chars)
    return {
        "text": text,
        "tables": [],
        "images": images,
        "charts": [],
        "metadata": {
            "page_count": len(text_parts),
            "image_count": len(images),
            "chart_count": 0,
            "text_chars": len(text),
        },
        "warnings": warnings,
    }


def extract_image(path: Path) -> dict:
    warnings = []
    images = [
        {
            "name": path.name,
            "path": str(path),
            "bytes": int(path.stat().st_size),
        }
    ]
    try:
        from PIL import Image

        with Image.open(str(path)) as image:
            images[0].update({"width": int(image.width), "height": int(image.height), "mode": str(image.mode)})
    except Exception as exc:
        warnings.append(f"image metadata extraction failed: {exc}")

    text = ""
    try:
        from PIL import Image
        import pytesseract

        with Image.open(str(path)) as image:
            text = str(pytesseract.image_to_string(image) or "").strip()
    except Exception as exc:
        warnings.append(f"local OCR unavailable or failed: {exc}")

    return {
        "text": text,
        "tables": [],
        "images": images,
        "charts": [],
        "metadata": {
            "image_count": 1,
            "chart_count": 0,
            "text_chars": len(text),
        },
        "warnings": warnings,
    }


def extract_plain_text(path: Path, max_text_chars: int) -> dict:
    warnings = []
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        warnings.append(f"plain text read failed: {exc}")
        text = ""
    text = trim_text(text, max_text_chars)
    return {
        "text": text,
        "tables": [],
        "images": [],
        "charts": [],
        "metadata": {"text_chars": len(text), "image_count": 0, "chart_count": 0},
        "warnings": warnings,
    }


def build_report(request: dict) -> dict:
    path = Path(str(request.get("source_path") or ""))
    source_file_type = detect_type(path, str(request.get("file_type") or ""))
    max_text_chars = int(request.get("max_text_chars") or 60000)
    if not path.exists() or not path.is_file():
        return {"status": "failed", "failure_reason": "source file does not exist", "source_file_type": source_file_type}

    suffix = path.suffix.lower()
    if source_file_type in {"docx"} or suffix == ".docx":
        extracted = extract_docx(path, max_text_chars)
        source_file_type = "docx"
    elif source_file_type in {"pdf"} or suffix == ".pdf":
        extracted = extract_pdf(path, max_text_chars)
        source_file_type = "pdf"
    elif suffix in IMAGE_SUFFIXES or source_file_type in {item.lstrip(".") for item in IMAGE_SUFFIXES}:
        extracted = extract_image(path)
        source_file_type = source_file_type or suffix.lstrip(".")
    else:
        extracted = extract_plain_text(path, max_text_chars)

    metadata = dict(extracted.get("metadata") or {})
    metadata.update(
        {
            "source_size_bytes": int(path.stat().st_size),
            "source_suffix": suffix,
        }
    )
    return {
        "status": "success",
        "source_file": {
            "name": str(request.get("file_name") or path.name),
            "path": str(path),
            "type": source_file_type,
            "suffix": suffix,
            "size": int(path.stat().st_size),
        },
        "source_file_type": source_file_type,
        "text": str(extracted.get("text") or ""),
        "tables": list(extracted.get("tables") or []),
        "images": list(extracted.get("images") or []),
        "charts": list(extracted.get("charts") or []),
        "metadata": metadata,
        "warnings": list(extracted.get("warnings") or []),
    }


def main() -> None:
    request_path = Path(sys.argv[1])
    report_path = Path(sys.argv[2])
    request = json.loads(request_path.read_text(encoding="utf-8"))
    report = build_report(request)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("file extract completed")


if __name__ == "__main__":
    main()
'''.lstrip()

    def _success_file_extract_result(
        self,
        *,
        user: User,
        extraction_report: dict[str, Any],
        original_message: str,
    ) -> dict[str, Any]:
        source_file = extraction_report.get("source_file") if isinstance(extraction_report.get("source_file"), dict) else {}
        metadata = extraction_report.get("metadata") if isinstance(extraction_report.get("metadata"), dict) else {}
        summary = "File extraction completed successfully using local CodeAgent parser."
        tool_output = {
            "tool": "code_agent",
            "title": "Code Agent File Extract",
            "summary": summary,
            "data": {
                "status": "success",
                "file_extraction_report": extraction_report,
                "original_message": original_message,
            },
            "card": {
                "type": "resume_card",
                "tool": "code_agent",
                "title": "Code Agent File Extract",
                "summary": summary,
                "data": {
                    "status": "success",
                    "source_file": source_file,
                    "metadata": metadata,
                },
            },
            "next_actions": [],
            "context_patch": {},
        }
        return {
            "status": "success",
            "reply": summary,
            "tool_steps": [
                {"tool": "code_file_extract_generate", "status": "done", "text": "done: generated controlled extractor.py"},
                {"tool": "code_file_extract_run", "status": "done", "text": "done: local file extraction completed"},
            ],
            "artifacts": [],
            "context_patch": {},
            "requires_user_input": False,
            "question": "",
            "code_task": {
                "type": "file_extract",
                "status": "success",
                "language": "python",
                "passed": True,
                "student_user_id": int(getattr(user, "id", 0) or 0),
            },
            "tool_outputs": [tool_output],
            "verification": extraction_report,
            "tests": [],
            "files": [{"path": "extractor.py"}, {"path": "report.json"}],
        }

    @staticmethod
    def _failed_file_extract_result(
        *,
        failure_reason: str,
        extraction_report: dict[str, Any],
        original_message: str,
    ) -> dict[str, Any]:
        reason = str(failure_reason or "File extraction failed.")
        tool_output = {
            "tool": "code_agent",
            "title": "Code Agent File Extract",
            "summary": "File extraction failed strict verification.",
            "data": {
                "status": "failed",
                "file_extraction_report": extraction_report,
                "failure_reason": reason,
                "original_message": original_message,
            },
            "card": {
                "type": "resume_card",
                "tool": "code_agent",
                "title": "Code Agent File Extract",
                "summary": "File extraction failed strict verification.",
                "data": {"status": "failed", "failure_reason": reason},
            },
            "next_actions": [],
            "context_patch": {},
        }
        return {
            "status": "failed",
            "reply": "File extraction failed strict verification.",
            "tool_steps": [
                {"tool": "code_file_extract", "status": "failed", "text": f"failed: {reason}"},
            ],
            "artifacts": [],
            "context_patch": {},
            "requires_user_input": False,
            "question": reason,
            "code_task": {
                "type": "file_extract",
                "status": "failed",
                "language": "python",
                "passed": False,
                "failure_reason": reason,
            },
            "tool_outputs": [tool_output],
            "verification": extraction_report,
            "tests": [],
            "files": [],
        }

    def _success_document_render_result(
        self,
        *,
        user: User,
        render_report: dict[str, Any],
        artifact: dict[str, Any],
        original_message: str,
    ) -> dict[str, Any]:
        summary = "Document render completed successfully. DOCX passed render tests."
        tool_output = {
            "tool": "code_agent",
            "title": "Code Agent Document Render",
            "summary": summary,
            "data": {
                "status": "success",
                "render_report": render_report,
                "original_message": original_message,
            },
            "card": {
                "type": "action_checklist_card",
                "tool": "code_agent",
                "title": "Code Agent Document Render",
                "summary": summary,
                "data": {"status": "success", "artifact": artifact},
            },
            "next_actions": [],
            "context_patch": {},
        }
        return {
            "status": "success",
            "reply": summary,
            "tool_steps": [
                {"tool": "code_document_render_generate", "status": "done", "text": "done: generated main.py and tests/test_render.py"},
                {"tool": "code_document_render_test", "status": "done", "text": "done: DOCX render tests passed"},
                {"tool": "code_document_render_output", "status": "done", "text": "done: DOCX artifact created"},
            ],
            "artifacts": [artifact],
            "context_patch": {},
            "requires_user_input": False,
            "question": "",
            "code_task": {
                "type": "document_render",
                "status": "success",
                "language": "python",
                "passed": True,
                "student_user_id": int(getattr(user, "id", 0) or 0),
            },
            "tool_outputs": [tool_output],
            "verification": render_report,
            "tests": [{"path": "tests/test_render.py"}],
            "files": [{"path": "main.py"}, {"path": "tests/test_render.py"}],
        }

    @staticmethod
    def _failed_document_render_result(
        *,
        failure_reason: str,
        render_report: dict[str, Any],
        original_message: str,
    ) -> dict[str, Any]:
        reason = str(failure_reason or "Document render failed.")
        tool_output = {
            "tool": "code_agent",
            "title": "Code Agent Document Render",
            "summary": "Document render failed strict verification.",
            "data": {
                "status": "failed",
                "render_report": render_report,
                "failure_reason": reason,
                "original_message": original_message,
            },
            "card": {
                "type": "action_checklist_card",
                "tool": "code_agent",
                "title": "Code Agent Document Render",
                "summary": "Document render failed strict verification.",
                "data": {"status": "failed", "failure_reason": reason},
            },
            "next_actions": [],
            "context_patch": {},
        }
        return {
            "status": "failed",
            "reply": "Document render failed strict verification.",
            "tool_steps": [
                {"tool": "code_document_render", "status": "failed", "text": f"failed: {reason}"},
            ],
            "artifacts": [],
            "context_patch": {},
            "requires_user_input": False,
            "question": reason,
            "code_task": {
                "type": "document_render",
                "status": "failed",
                "language": "python",
                "passed": False,
                "failure_reason": reason,
            },
            "tool_outputs": [tool_output],
            "verification": render_report,
            "tests": [],
            "files": [],
        }

    def _detect_language(self, lowered_text: str) -> str:
        text = str(lowered_text or "")
        for language, keywords in self.LANGUAGE_KEYWORDS:
            if any(keyword in text for keyword in keywords):
                return language
        return ""

    def _generate_code_package(
        self,
        *,
        user: User,
        message: str,
        language: str,
        attempt: int,
        constraints: dict[str, Any],
        previous_errors: list[str],
    ) -> dict[str, Any]:
        prompt = self._build_generation_prompt(
            language=language,
            message=message,
            attempt=attempt,
            constraints=constraints,
            previous_errors=previous_errors,
        )
        raw = self._request_model(user=user, prompt=prompt)
        parsed = CodeExecutionRunner.parse_model_output(raw)
        normalized = self._normalize_package(parsed=parsed, language=language)
        return {
            "source": "llm",
            "files": normalized.get("files") or [],
            "tests": normalized.get("tests") or [],
        }

    def _build_generation_prompt(
        self,
        *,
        language: str,
        message: str,
        attempt: int,
        constraints: dict[str, Any],
        previous_errors: list[str],
    ) -> str:
        previous = "\n".join(f"- {item}" for item in previous_errors if item) or "- None"
        constraints_json = json.dumps(constraints, ensure_ascii=False)
        return (
            "You are a strict Code Agent.\n"
            "Your task is to generate complete code and self-test scripts.\n"
            "Return ONLY one JSON object, no markdown.\n"
            "JSON schema:\n"
            "{\n"
            "  \"language\": \"python|c|cpp|javascript|html|css|vue|vbs|mermaid\",\n"
            "  \"entry_file\": \"relative path\",\n"
            "  \"files\": [{\"path\": \"relative path\", \"content\": \"file content\"}],\n"
            "  \"tests\": [{\"path\": \"tests/test_x.ext\", \"content\": \"self-test script\"}]\n"
            "}\n"
            "Rules:\n"
            "1. No network usage in any file.\n"
            "2. Keep code runnable without external secrets.\n"
            "3. For Python, include at least one assert-based test script.\n"
            "4. Paths must be relative, no absolute path, no ../ traversal.\n"
            f"Target language: {language}\n"
            f"Attempt: {attempt}\n"
            f"User requirement: {message}\n"
            f"Detected constraints: {constraints_json}\n"
            "Previous verification errors:\n"
            f"{previous}\n"
            "Output JSON now:"
        )

    def _request_model(self, *, user: User, prompt: str) -> str:
        role = "student"
        if getattr(user, "role", None) and getattr(user.role, "code", None):
            role = str(user.role.code)
        user_name = str(getattr(user, "real_name", "Code User") or "Code User")
        return str(
            self.llm_service.chat(
                user_role=role,
                user_name=user_name,
                message=prompt,
                history=[],
                context={
                    "scene": "code_generation",
                    "selected_skill": "code-agent",
                    "intent": "code_generation",
                    "reply_mode": "structured",
                    "tool_outputs": [],
                    "retrieval_chunks": [],
                    "slots": {},
                    "small_talk": False,
                },
            )
            or ""
        )

    def _normalize_package(self, *, parsed: dict[str, Any], language: str) -> dict[str, Any]:
        raw_files = parsed.get("files") if isinstance(parsed.get("files"), list) else []
        raw_tests = parsed.get("tests") if isinstance(parsed.get("tests"), list) else []

        if not raw_files:
            return {"from_llm": True, "files": [], "tests": []}

        files = self._normalize_file_list(raw_files, default_name=self.DEFAULT_FILE_NAME.get(language, "main.txt"))
        tests = self._normalize_test_list(raw_tests)

        if not files:
            return {"from_llm": True, "files": [], "tests": tests}

        return {
            "from_llm": True,
            "files": files,
            "tests": tests,
        }

    def _normalize_file_list(self, rows: list[dict[str, Any]], *, default_name: str) -> list[dict[str, str]]:
        files: list[dict[str, str]] = []
        for index, row in enumerate(rows):
            path_raw = str((row or {}).get("path") or "").strip()
            content = str((row or {}).get("content") or "")
            if not content.strip():
                continue
            safe_path = self._safe_relative_path(path_raw)
            if not safe_path:
                safe_path = default_name if index == 0 else f"file_{index}_{Path(default_name).name}"
            files.append({"path": safe_path, "content": content})
        return files

    def _normalize_test_list(self, rows: list[dict[str, Any]]) -> list[dict[str, str]]:
        tests: list[dict[str, str]] = []
        for index, row in enumerate(rows):
            path_raw = str((row or {}).get("path") or "").strip()
            content = str((row or {}).get("content") or "")
            if not content.strip():
                continue
            safe_path = self._safe_relative_path(path_raw)
            if not safe_path:
                suffix = Path(path_raw).suffix or ".txt"
                safe_path = f"tests/test_{index}{suffix}"
            if not safe_path.startswith("tests/"):
                safe_path = f"tests/{Path(safe_path).name}"
            tests.append({"path": safe_path, "content": content})
        return tests

    @staticmethod
    def _safe_relative_path(path_text: str) -> str:
        text = str(path_text or "").strip().replace("\\", "/")
        if not text or text.startswith("/") or ":" in text:
            return ""
        parts = text.split("/")
        if any(part in {"", ".."} for part in parts):
            return ""
        return text

    def _success_result(
        self,
        *,
        user: User,
        run_id: str,
        language: str,
        attempt: int,
        files: list[dict[str, str]],
        tests: list[dict[str, str]],
        verification_report: dict[str, Any],
        attempts: list[dict[str, Any]],
        constraints: dict[str, Any],
        original_message: str,
    ) -> dict[str, Any]:
        exported = self._export_files(user=user, run_id=run_id, files=files)
        artifacts = exported["artifacts"]
        bundle_files = exported["bundle_files"]

        bundle = {
            "run_id": run_id,
            "status": "passed",
            "language": language,
            "attempt": attempt,
            "max_attempts": self.MAX_ATTEMPTS,
            "verification_report": verification_report,
            "files": bundle_files,
            "tests": tests,
            "constraints": constraints,
            "original_message": original_message,
            "artifacts": artifacts,
            "model": {
                "provider": str(getattr(self.llm_service, "provider", "unknown") or "unknown"),
                "model_name": str(getattr(self.llm_service, "model_name", "unknown") or "unknown"),
            },
            "attempt_reports": attempts,
            "bundle_meta": {
                "file_count": len(bundle_files),
                "test_count": len(tests),
                "artifact_count": len(artifacts),
            },
        }

        summary = "Code generated successfully. All required compile/syntax and self-test checks passed."
        tool_steps = [
            {"tool": "code_generate", "status": "done", "text": f"done: generated code in attempt {attempt}"},
            {"tool": "code_verify", "status": "done", "text": "done: compile/syntax checks passed"},
            {"tool": "code_test", "status": "done", "text": "done: self-tests passed"},
        ]
        context_patch = {
            "context_binding": {
                "generated_code_bundle": bundle,
                "pending_file_offer": {
                    "source": "code_agent",
                    "task_type": "generate_document",
                    "run_id": run_id,
                    "language": language,
                },
            }
        }
        tool_output = {
            "tool": "code_agent",
            "title": "Code Agent",
            "summary": summary,
            "data": {
                "status": "success",
                "language": language,
                "attempt": attempt,
                "files": bundle_files,
                "verification_report": verification_report,
                "attempt_reports": attempts,
                "constraints": constraints,
                "bundle_meta": bundle.get("bundle_meta") or {},
            },
            "card": {
                "type": "action_checklist_card",
                "tool": "code_agent",
                "title": "Code Agent",
                "summary": summary,
                "data": {
                    "status": "success",
                    "language": language,
                    "artifact_count": len(artifacts),
                    "file_count": len(bundle_files),
                },
            },
            "next_actions": [
                "Refine code style",
                "Add more edge-case tests",
                "Generate another language version",
            ],
            "context_patch": context_patch,
        }

        return {
            "status": "success",
            "reply": summary,
            "tool_steps": tool_steps,
            "artifacts": artifacts,
            "context_patch": context_patch,
            "requires_user_input": False,
            "question": "",
            "code_task": {
                "type": "code_generation",
                "status": "success",
                "language": language,
                "attempt": attempt,
                "max_attempts": self.MAX_ATTEMPTS,
                "passed": True,
            },
            "tool_outputs": [tool_output],
            "verification": verification_report,
            "tests": tests,
            "files": bundle_files,
        }

    def _failed_result(self, *, language: str, attempts: list[dict[str, Any]], message: str, constraints: dict[str, Any]) -> dict[str, Any]:
        last_report = attempts[-1]["verification_report"] if attempts else {}
        failure_summary = CodeExecutionRunner.build_verification_summary(last_report) if last_report else "Code generation failed."
        suggestions = self._build_retry_suggestions(language=language, original_message=message, verification=last_report)
        failure_explanation = self._build_failure_explanation(verification=last_report, attempts=attempts)
        summary = (
            "Code generation was blocked because required verification did not fully pass. "
            "No final code was output."
        )
        tool_steps = [
            {"tool": "code_generate", "status": "failed", "text": f"failed: reached max attempts {self.MAX_ATTEMPTS}"},
            {"tool": "code_verify", "status": "failed", "text": f"failed: {failure_summary}"},
            {"tool": "code_test", "status": "failed", "text": "failed: strict output gate triggered"},
        ]
        tool_output = {
            "tool": "code_agent",
            "title": "Code Agent",
            "summary": summary,
            "data": {
                "status": "failed",
                "language": language,
                "attempt": len(attempts),
                "verification_report": last_report,
                "attempt_reports": attempts,
                "fix_suggestions": suggestions,
                "failure_explanation": failure_explanation,
                "constraints": constraints,
            },
            "card": {
                "type": "action_checklist_card",
                "tool": "code_agent",
                "title": "Code Agent",
                "summary": summary,
                "data": {
                    "status": "failed",
                    "language": language,
                    "attempts": len(attempts),
                },
            },
            "next_actions": suggestions,
            "context_patch": {},
        }

        return {
            "status": "failed",
            "reply": summary,
            "tool_steps": tool_steps,
            "artifacts": [],
            "context_patch": {},
            "requires_user_input": False,
            "question": f"Verification failed: {failure_summary}",
            "code_task": {
                "type": "code_generation",
                "status": "failed",
                "language": language,
                "attempt": len(attempts),
                "max_attempts": self.MAX_ATTEMPTS,
                "passed": False,
                "failure_reason": failure_summary,
                "failure_explanation": failure_explanation,
            },
            "tool_outputs": [tool_output],
            "verification": last_report,
            "tests": [],
            "files": [],
        }

    def _export_files(self, *, user: User, run_id: str, files: list[dict[str, str]]) -> dict[str, Any]:
        user_id = int(getattr(user, "id", 0) or 0)
        folder = self.settings.upload_path / "agent_code" / f"user_{user_id}" / run_id
        folder.mkdir(parents=True, exist_ok=True)

        artifacts: list[dict[str, Any]] = []
        bundle_files: list[dict[str, Any]] = []

        for row in files:
            safe_path = self._safe_relative_path(row["path"])
            if not safe_path:
                continue
            target = folder / safe_path
            target.parent.mkdir(parents=True, exist_ok=True)
            content = row["content"]
            target.write_text(content, encoding="utf-8")
            digest = hashlib.sha256(content.encode("utf-8")).hexdigest()
            download_url = self._to_upload_url(target)

            artifact = {
                "name": Path(safe_path).name,
                "type": "code",
                "language": Path(safe_path).suffix.lstrip(".").lower(),
                "download_url": download_url,
                "mime_type": self._mime_by_suffix(Path(safe_path).suffix.lower()),
            }
            bundle_entry = {
                "path": safe_path,
                "content": content,
                "size": len(content.encode("utf-8")),
                "sha256": digest,
                "download_url": download_url,
            }
            artifacts.append(artifact)
            bundle_files.append(bundle_entry)

        return {"artifacts": artifacts, "bundle_files": bundle_files}

    def _to_upload_url(self, file_path: Path) -> str:
        relative = file_path.resolve().relative_to(self.settings.upload_path.resolve()).as_posix()
        return f"/uploads/{relative}"

    @staticmethod
    def _mime_by_suffix(suffix: str) -> str:
        mapping = {
            ".py": "text/x-python",
            ".c": "text/x-c",
            ".cpp": "text/x-c++src",
            ".cc": "text/x-c++src",
            ".cxx": "text/x-c++src",
            ".js": "text/javascript",
            ".html": "text/html",
            ".htm": "text/html",
            ".css": "text/css",
            ".vue": "text/plain",
            ".vbs": "text/vbscript",
            ".mmd": "text/plain",
            ".mermaid": "text/plain",
            ".md": "text/markdown",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        return mapping.get(suffix, "text/plain")

    def _extract_constraints(self, message: str) -> dict[str, Any]:
        text = str(message or "")
        lowered = text.lower()
        max_files_match = re.search(r"(?:最多|不超过|max)\s*(\d+)\s*(?:个)?(?:文件|files?)", lowered)
        runtime_match = re.search(r"(python\s*\d+(?:\.\d+)?)", lowered)
        output_match = re.search(r"(json|yaml|csv|markdown|md|html|pdf)", lowered)
        return {
            "need_multi_file": any(token in lowered for token in ("multi-file", "多文件", "多个文件")),
            "need_example_io": any(token in lowered for token in ("example", "示例输入", "示例输出", "样例")),
            "prefer_class": any(token in lowered for token in ("class", "类")),
            "prefer_functional": any(token in lowered for token in ("function", "函数")),
            "forbid_network": any(token in lowered for token in ("离线", "不能联网", "no network", "offline")),
            "prefer_cli": any(token in lowered for token in ("命令行", "cli", "terminal")),
            "need_readme": any(token in lowered for token in ("readme", "说明文档", "使用说明")),
            "max_files": int(max_files_match.group(1)) if max_files_match else 0,
            "runtime_hint": runtime_match.group(1) if runtime_match else "",
            "output_format": output_match.group(1) if output_match else "",
        }

    def _build_failure_explanation(self, *, verification: dict[str, Any], attempts: list[dict[str, Any]]) -> dict[str, Any]:
        compile_block = verification.get("compile") if isinstance(verification.get("compile"), dict) else {}
        test_block = verification.get("tests") if isinstance(verification.get("tests"), dict) else {}
        return {
            "attempt_count": len(attempts),
            "compile_summary": str(compile_block.get("summary") or ""),
            "compile_errors": list(compile_block.get("errors") or [])[:5],
            "test_summary": str(test_block.get("summary") or ""),
            "test_errors": list(test_block.get("errors") or [])[:5],
        }

    def _build_retry_suggestions(self, *, language: str, original_message: str, verification: dict[str, Any]) -> list[str]:
        summary = CodeExecutionRunner.build_verification_summary(verification)
        suggestions = [
            f"Clarify expected runtime behavior for {language}.",
            "Provide stricter input/output examples.",
            "Request smaller scoped implementation first.",
        ]
        if "Network capability detected" in summary:
            suggestions.insert(0, "Remove all network-related dependencies from the requirement.")
        if "self-test script" in summary:
            suggestions.insert(0, "Ask the model to return explicit assert-based test scripts.")
        if not str(original_message or "").strip():
            suggestions.insert(0, "Describe the expected feature in one sentence before regenerating.")
        return suggestions[:4]

