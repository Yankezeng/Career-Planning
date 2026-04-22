from __future__ import annotations

import json
import subprocess
import sys

from docx import Document

from app.services.agent.code_agent.code_agent import CodeAgent
from app.services.assistant_background_jobs import _friendly_failure_reason


def _minimal_resume_spec() -> dict:
    return {
        "target_role": "Java后端开发工程师",
        "optimized_resume_document": "\n".join(
            [
                "李悦",
                "求职意向：Java后端开发工程师",
                "电话：13598761234 | 邮箱：liyue@example.com",
                "",
                "教育背景",
                "贵州大学 | 计算机科学与技术 | 本科 | 2022.09 - 2026.06",
                "- GPA：3.7/4.0，专业排名前15%",
                "",
                "专业技能",
                "- 熟悉 Java、Spring Boot、MySQL、Redis 和接口开发。",
                "- 熟悉 Git、Linux 基础命令和常见后端工程化流程。",
                "",
                "项目经历",
                "校园招聘系统 | Java后端开发 | 2025.01 - 2025.03",
                "- 使用 Spring Boot 完成岗位、简历和投递接口开发。",
                "- 通过索引优化将常用查询响应时间降低约 30%。",
            ]
        ),
    }


def _document_text(path) -> str:
    document = Document(str(path))
    chunks = [paragraph.text for paragraph in document.paragraphs]
    chunks.extend(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return "\n".join(chunks)


def test_document_renderer_template_imports_json() -> None:
    renderer_code = CodeAgent._build_document_renderer_code()

    assert "\nimport json\n" in renderer_code
    assert "spec = json.loads" in renderer_code


def test_document_renderer_cli_generates_docx(tmp_path) -> None:
    main_path = tmp_path / "main.py"
    spec_path = tmp_path / "spec.json"
    output_path = tmp_path / "optimized-resume.docx"
    main_path.write_text(CodeAgent._build_document_renderer_code(), encoding="utf-8")
    spec_path.write_text(json.dumps(_minimal_resume_spec(), ensure_ascii=False, indent=2), encoding="utf-8")

    completed = subprocess.run(
        [sys.executable, str(main_path), str(spec_path), str(output_path)],
        cwd=str(tmp_path),
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr or completed.stdout
    assert output_path.exists()
    assert output_path.stat().st_size > 5 * 1024
    text = _document_text(output_path)
    assert "Java后端开发工程师" in text
    assert "教育背景" in text
    assert "专业技能" in text


def test_background_failure_reason_hides_traceback() -> None:
    reason = _friendly_failure_reason(
        RuntimeError(
            'DOCX render command failed: Traceback (most recent call last):\n'
            '  File "main.py", line 468, in main\n'
            "    spec = json.loads(...)\n"
            "NameError: name 'json' is not defined"
        )
    )

    assert "Word 文档生成失败" in reason
    assert "Traceback" not in reason
    assert "NameError" not in reason
    assert "json" not in reason
